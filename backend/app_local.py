"""
Local development version of the RAG Chatbot PWA
Simplified setup without Docker, Redis, or complex dependencies

This module provides a complete Flask application for local development with:
- SQLite database for data storage
- JWT authentication for secure API access
- CORS enabled for frontend integration
- Comprehensive logging and monitoring
- File upload and document processing
- Vector search with FAISS
- RAG (Retrieval-Augmented Generation) chat functionality
- Admin dashboard for system management

Features:
- Multi-format document processing (PDF, DOCX, TXT, etc.)
- Context-based knowledge management
- Real-time chat with streaming responses
- User authentication and authorization
- Security middleware with rate limiting
- Comprehensive error handling and logging
- Health checks and system monitoring

Usage:
    python app_local.py

Environment Variables:
    JWT_SECRET_KEY: Secret key for JWT token generation
    DATABASE_URL: Database connection string (default: SQLite)
    LOG_LEVEL: Logging level (DEBUG, INFO, WARNING, ERROR)
    GEMINI_API_KEY: Google Gemini API key for LLM functionality
    UPLOAD_FOLDER: Directory for file uploads (default: uploads)
    MAX_CONTENT_LENGTH: Maximum file upload size (default: 16MB)

Author: RAG Chatbot Development Team
Version: 1.0.0
"""

import os
from datetime import datetime, timedelta, timezone
from flask import Flask, request, jsonify, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import json
import re
from pathlib import Path
import mimetypes

# Load environment variables
from dotenv import load_dotenv

# Setup logging first
from logging_config import setup_logging, get_logger, log_request_info, log_error_with_context

# Initialize logging
loggers = setup_logging(
    log_level=os.getenv('LOG_LEVEL', 'INFO'),
    console_output=True,
    json_format=False
)

# Get application logger
app_logger = get_logger('app')

# Import security middleware
from security_middleware import SecurityMiddleware
load_dotenv()

# Create Flask app
app = Flask(__name__)
app.url_map.strict_slashes = False

# Configuration
import secrets
default_secret = secrets.token_urlsafe(32)
app.config['SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', default_secret)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///ragchatbot.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', default_secret)
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=24)
app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', '104857600'))  # 100MB

# Import database and models
from database import db
from models import User, Context, Document, ChatSession, Message, TextChunk

# Initialize the db with the app
db.init_app(app)

# Initialize other extensions
jwt = JWTManager(app)

# COMPREHENSIVE CORS CONFIGURATION for LOCAL DEVELOPMENT
# Allow ALL requests from localhost and 127.0.0.1 on ANY port
CORS(app, 
     origins=["*"],  # Allow all origins for development
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
     allow_headers=["Content-Type", "Authorization", "X-Requested-With", "Accept", "Origin"],
     supports_credentials=False,
     send_wildcard=True)

# Add manual CORS headers as backup for complex requests
@app.after_request
def after_request_cors(response):
    origin = request.headers.get('Origin')
    if origin and ('localhost' in origin or '127.0.0.1' in origin):
        response.headers['Access-Control-Allow-Origin'] = origin
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS, PATCH'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization, X-Requested-With, Accept, Origin'
        response.headers['Access-Control-Max-Age'] = '86400'
    return response

# Handle preflight OPTIONS requests
@app.before_request
def handle_options():
    if request.method == "OPTIONS":
        origin = request.headers.get('Origin')
        if origin and ('localhost' in origin or '127.0.0.1' in origin):
            from flask import make_response
            response = make_response()
            response.headers['Access-Control-Allow-Origin'] = origin
            response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS, PATCH'
            response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization, X-Requested-With, Accept, Origin'
            response.headers['Access-Control-Max-Age'] = '86400'
            return response

# DISABLE SECURITY MIDDLEWARE FOR LOCAL DEVELOPMENT
# Comment out the security middleware to avoid rate limiting during development
# security_middleware = SecurityMiddleware(app)
app_logger.info("Flask application initialized WITHOUT security middleware for local development")

# Create upload folder
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Add a simple CORS test endpoint
@app.route('/api/cors-test', methods=['GET', 'POST', 'OPTIONS'])
def cors_test():
    """Test endpoint to verify CORS is working"""
    if request.method == 'OPTIONS':
        from flask import make_response
        response = make_response()
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,X-Requested-With,Accept,Origin')
        response.headers.add('Access-Control-Allow-Methods', 'GET,POST,PUT,DELETE,OPTIONS,PATCH')
        return response
    return jsonify({'message': 'CORS is working!', 'method': request.method})

# Helper function to get user ID as integer from JWT
def get_current_user_id():
    """
    Get the current user ID as an integer from JWT token
    
    This function extracts the user identity from the JWT token in the current request
    context and converts it to an integer for database operations.
    
    Returns:
        int: The current user's ID if authenticated, None otherwise
        
    Raises:
        ValueError: If the JWT identity cannot be converted to an integer
        
    Example:
        >>> user_id = get_current_user_id()
        >>> if user_id:
        >>>     user = User.query.get(user_id)
    """
    try:
        user_id_str = get_jwt_identity()
        if user_id_str:
            user_id = int(user_id_str)
            app_logger.debug(f"Retrieved user ID {user_id} from JWT token")
            return user_id
        else:
            app_logger.debug("No user identity found in JWT token")
            return None
    except (ValueError, TypeError) as e:
        app_logger.error(f"Error converting JWT identity to integer: {e}")
        return None

# Text processing functions
def extract_text_from_file(file_path):
    """
    Extract text from various file types with language-specific handling
    
    This function is the main entry point for text extraction from uploaded files.
    It supports multiple file formats including documents, code files, and data files.
    Each file type is processed with specialized handlers to preserve structure and meaning.
    
    Supported File Types:
        - Documents: PDF (.pdf), Word (.docx, .doc), Excel (.xlsx, .xls)
        - Programming: Python (.py), JavaScript/TypeScript (.js, .ts), Java (.java),
          C/C++ (.cpp, .c, .h, .hpp), Go (.go), Rust (.rs)
        - Markup: Markdown (.md), HTML (.html), CSS (.css)
        - Data: JSON (.json), YAML (.yaml, .yml), CSV (.csv)
        - Text: Plain text (.txt) and other text-based files
    
    Args:
        file_path (str|Path): Path to the file to extract text from
        
    Returns:
        str: Extracted and formatted text content with metadata headers
        
    Raises:
        FileNotFoundError: If the specified file doesn't exist
        PermissionError: If unable to read the file
        UnicodeDecodeError: If file encoding cannot be determined
        
    Example:
        >>> text = extract_text_from_file('/path/to/document.pdf')
        >>> print(text[:100])  # First 100 characters
        
    Note:
        Binary files (PDF, DOCX, XLSX) require additional dependencies:
        - PyMuPDF for PDF processing
        - openpyxl/pandas for Excel files
        - python-docx for Word documents
    """
    file_path = Path(file_path)
    file_extension = file_path.suffix.lower()
    file_name = file_path.name
    
    app_logger.info(f"Starting text extraction for file: {file_name} (type: {file_extension})")
    
    try:
        # Handle binary file types first
        if file_extension == '.pdf':
            app_logger.debug(f"Processing PDF file: {file_name}")
            result = extract_pdf_content(file_path, file_name)
            app_logger.info(f"Successfully extracted {len(result)} characters from PDF: {file_name}")
            return result
        elif file_extension in ['.xlsx', '.xls']:
            app_logger.debug(f"Processing Excel file: {file_name}")
            result = extract_excel_content(file_path, file_name)
            app_logger.info(f"Successfully extracted {len(result)} characters from Excel: {file_name}")
            return result
        elif file_extension in ['.docx', '.doc']:
            app_logger.debug(f"Processing Word document: {file_name}")
            result = extract_docx_content(file_path, file_name)
            app_logger.info(f"Successfully extracted {len(result)} characters from Word: {file_name}")
            return result

        # Handle text-based files
        app_logger.debug(f"Processing text-based file: {file_name}")
        content = None
        encoding_used = 'utf-8'
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            app_logger.debug(f"Successfully read file with UTF-8 encoding")
        except UnicodeDecodeError as e:
            app_logger.warning(f"UTF-8 decoding failed for {file_name}, trying latin-1: {e}")
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                encoding_used = 'latin-1'
                app_logger.debug(f"Successfully read file with latin-1 encoding")
            except Exception as e2:
                app_logger.error(f"Failed to read file with latin-1 encoding: {e2}")
                raise e2

        if content is None:
            raise ValueError("Failed to read file content")

        # Add language-specific metadata and formatting
        app_logger.debug(f"Applying language-specific formatting for {file_extension}")
        
        if file_extension in ['.py']:
            formatted_content = format_python_content(content, file_name)
        elif file_extension in ['.js', '.ts']:
            formatted_content = format_javascript_content(content, file_name)
        elif file_extension in ['.java']:
            formatted_content = format_java_content(content, file_name)
        elif file_extension in ['.cpp', '.c', '.h', '.hpp']:
            formatted_content = format_cpp_content(content, file_name)
        elif file_extension in ['.go']:
            formatted_content = format_go_content(content, file_name)
        elif file_extension in ['.rs']:
            formatted_content = format_rust_content(content, file_name)
        elif file_extension == '.md':
            formatted_content = format_markdown_content(content, file_name)
        elif file_extension in ['.json', '.yaml', '.yml']:
            formatted_content = format_config_content(content, file_name, file_extension)
        elif file_extension == '.csv':
            formatted_content = format_csv_content(content, file_name)
        elif file_extension == '.txt':
            formatted_content = f"# {file_name}\n\n{content}"
        else:
            # For other file types, add basic formatting
            formatted_content = f"# {file_name}\nFile Type: {file_extension}\nEncoding: {encoding_used}\n\n{content}"

        app_logger.info(f"Successfully extracted and formatted {len(formatted_content)} characters from {file_name}")
        return formatted_content

    except FileNotFoundError as e:
        error_msg = f"File not found: {file_path}"
        app_logger.error(error_msg)
        log_error_with_context(e, {"file_path": str(file_path), "file_name": file_name})
        return f"# {file_name}\nError: {error_msg}"
    except PermissionError as e:
        error_msg = f"Permission denied accessing file: {file_path}"
        app_logger.error(error_msg)
        log_error_with_context(e, {"file_path": str(file_path), "file_name": file_name})
        return f"# {file_name}\nError: {error_msg}"
    except Exception as e:
        error_msg = f"Unexpected error reading file {file_path}: {str(e)}"
        app_logger.error(error_msg)
        log_error_with_context(e, {"file_path": str(file_path), "file_name": file_name, "file_extension": file_extension})
        return f"# {file_name}\nError: {error_msg}"

def extract_pdf_content(file_path, file_name):
    """
    Extract text from PDF files using PyMuPDF (fitz)
    
    This function extracts text content from PDF documents page by page,
    preserving page structure and handling various PDF formats including
    text-based PDFs and OCR-scanned documents.
    
    Args:
        file_path (str|Path): Path to the PDF file
        file_name (str): Name of the file for header generation
        
    Returns:
        str: Formatted text content with page headers and metadata
        
    Raises:
        ImportError: If PyMuPDF is not installed
        Exception: For PDF parsing errors or corrupted files
        
    Dependencies:
        - PyMuPDF (fitz): pip install pymupdf
        
    Example:
        >>> content = extract_pdf_content('/path/to/document.pdf', 'document.pdf')
        >>> print(f"Extracted {len(content)} characters from PDF")
        
    Note:
        - Empty pages are automatically skipped
        - Page numbers are preserved in the output
        - Metadata about the document is included in headers
    """
    app_logger.debug(f"Starting PDF text extraction for: {file_name}")
    
    try:
        import fitz  # PyMuPDF
        app_logger.debug("PyMuPDF imported successfully")

        # Open PDF document
        doc = fitz.open(str(file_path))
        text_content = []
        total_pages = len(doc)
        pages_with_content = 0
        
        app_logger.info(f"Processing PDF with {total_pages} pages: {file_name}")

        # Extract text from each page
        for page_num in range(total_pages):
            try:
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    text_content.append(f"## Page {page_num + 1}\n\n{text}")
                    pages_with_content += 1
                    app_logger.debug(f"Extracted {len(text)} characters from page {page_num + 1}")
                else:
                    app_logger.debug(f"Page {page_num + 1} is empty or contains no text")
                    
            except Exception as page_error:
                app_logger.warning(f"Error processing page {page_num + 1}: {page_error}")
                text_content.append(f"## Page {page_num + 1}\n\n[Error extracting page content: {page_error}]")

        doc.close()
        app_logger.debug("PDF document closed successfully")

        # Create formatted content with metadata
        metadata = [
            f"# PDF Document: {file_name}",
            f"Total Pages: {total_pages}",
            f"Pages with Content: {pages_with_content}",
            f"Extraction Date: {datetime.now().isoformat()}",
            ""
        ]
        
        full_content = "\n".join(metadata) + "\n".join(text_content)
        
        app_logger.info(f"Successfully extracted PDF content: {len(full_content)} characters from {pages_with_content}/{total_pages} pages")
        return full_content

    except ImportError as e:
        error_msg = "PyMuPDF not installed. Cannot extract PDF content. Install with: pip install pymupdf"
        app_logger.error(f"PDF extraction failed for {file_name}: {error_msg}")
        log_error_with_context(e, {"file_name": file_name, "file_path": str(file_path)})
        return f"# PDF Document: {file_name}\n\nError: {error_msg}"
    except Exception as e:
        error_msg = f"Error extracting PDF content: {str(e)}"
        app_logger.error(f"PDF extraction failed for {file_name}: {error_msg}")
        log_error_with_context(e, {"file_name": file_name, "file_path": str(file_path)})
        return f"# PDF Document: {file_name}\n\nError: {error_msg}"

def extract_excel_content(file_path, file_name):
    """Extract content from Excel files with enhanced analysis"""
    try:
        import pandas as pd
        import openpyxl
        from openpyxl import load_workbook

        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        content_parts = [f"# Excel File: {file_name}\n"]
        
        # Try to get workbook metadata
        try:
            wb = load_workbook(file_path, read_only=True)
            if hasattr(wb, 'properties') and wb.properties:
                props = wb.properties
                if props.title:
                    content_parts.append(f"**Title:** {props.title}")
                if props.creator:
                    content_parts.append(f"**Author:** {props.creator}")
                if props.description:
                    content_parts.append(f"**Description:** {props.description}")
                content_parts.append("")
            wb.close()
        except:
            pass  # Continue without metadata if it fails

        total_rows = 0
        total_sheets = len(excel_file.sheet_names)
        content_parts.append(f"**Total Sheets:** {total_sheets}\n")

        for i, sheet_name in enumerate(excel_file.sheet_names):
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                total_rows += len(df)

                content_parts.append(f"\n## Sheet {i+1}: {sheet_name}")
                content_parts.append(f"- **Dimensions:** {len(df)} rows × {len(df.columns)} columns")
                
                if not df.empty:
                    # Column information
                    content_parts.append(f"- **Columns:** {', '.join(df.columns.tolist())}")
                    
                    # Data types
                    dtype_summary = df.dtypes.value_counts()
                    dtype_info = []
                    for dtype, count in dtype_summary.items():
                        dtype_info.append(f"{count} {dtype}")
                    content_parts.append(f"- **Data Types:** {', '.join(dtype_info)}")

                    # Missing data info
                    missing_data = df.isnull().sum()
                    if missing_data.sum() > 0:
                        missing_cols = [f"{col} ({missing_data[col]})" for col in missing_data.index if missing_data[col] > 0]
                        content_parts.append(f"- **Missing Values:** {', '.join(missing_cols)}")

                    # Sample data (first 3 rows in markdown table format)
                    content_parts.append("\n### Sample Data (First 3 Rows)")
                    sample_df = df.head(3)
                    
                    # Create markdown table
                    headers = ["| " + " | ".join(str(col) for col in sample_df.columns) + " |"]
                    separator = ["| " + " | ".join(["---"] * len(sample_df.columns)) + " |"]
                    
                    rows = []
                    for _, row in sample_df.iterrows():
                        row_data = []
                        for val in row:
                            if pd.isna(val):
                                row_data.append("*empty*")
                            else:
                                val_str = str(val)
                                # Truncate very long values
                                if len(val_str) > 50:
                                    val_str = val_str[:47] + "..."
                                row_data.append(val_str)
                        rows.append("| " + " | ".join(row_data) + " |")
                    
                    content_parts.extend(headers + separator + rows)

                    # Summary statistics for numeric columns
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        content_parts.append("\n### Numeric Summary")
                        desc = df[numeric_cols].describe()
                        
                        # Format as markdown table
                        stats_headers = ["| Statistic | " + " | ".join(numeric_cols) + " |"]
                        stats_separator = ["| --- | " + " | ".join(["---"] * len(numeric_cols)) + " |"]
                        
                        stats_rows = []
                        for stat in ['count', 'mean', 'std', 'min', '25%', '50%', '75%', 'max']:
                            if stat in desc.index:
                                row_data = [stat]
                                for col in numeric_cols:
                                    val = desc.loc[stat, col]
                                    if pd.isna(val):
                                        row_data.append("N/A")
                                    else:
                                        if isinstance(val, float):
                                            row_data.append(f"{val:.2f}")
                                        else:
                                            row_data.append(str(val))
                                stats_rows.append("| " + " | ".join(row_data) + " |")
                        
                        content_parts.extend(stats_headers + stats_separator + stats_rows)

                    # Categorical data summary
                    categorical_cols = df.select_dtypes(include=['object', 'category']).columns
                    if len(categorical_cols) > 0:
                        content_parts.append("\n### Categorical Data Summary")
                        for col in categorical_cols[:3]:  # Limit to first 3 categorical columns
                            unique_count = df[col].nunique()
                            content_parts.append(f"- **{col}:** {unique_count} unique values")
                            if unique_count <= 10:  # Show values if not too many
                                top_values = df[col].value_counts().head(5)
                                values_str = ", ".join([f"{val} ({count})" for val, count in top_values.items()])
                                content_parts.append(f"  - Most frequent: {values_str}")

            except Exception as e:
                content_parts.append(f"\n## Sheet {i+1}: {sheet_name}")
                content_parts.append(f"Error reading sheet: {str(e)}")

        # Overall summary
        content_parts.append(f"\n## Overall Summary")
        content_parts.append(f"- **Total Sheets:** {total_sheets}")
        content_parts.append(f"- **Total Rows:** {total_rows}")

        return "\n".join(content_parts)

    except ImportError as e:
        missing_module = "pandas" if "pandas" not in str(e) else "openpyxl"
        return f"# Excel File: {file_name}\n\nError: {missing_module} not installed. Install with: pip install {missing_module} pandas"
    except Exception as e:
        return f"# Excel File: {file_name}\n\nError extracting Excel content: {str(e)}"

def extract_docx_content(file_path, file_name):
    """Extract content from Word documents with enhanced structure"""
    try:
        from docx import Document

        doc = Document(file_path)
        content_parts = [f"# Word Document: {file_name}\n"]

        # Document properties
        props = doc.core_properties
        if props.title:
            content_parts.append(f"**Title:** {props.title}")
        if props.author:
            content_parts.append(f"**Author:** {props.author}")
        if props.subject:
            content_parts.append(f"**Subject:** {props.subject}")
        content_parts.append("")

        # Extract paragraphs with style information
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            style_name = para.style.name if para.style else "Normal"
            
            # Handle headings
            if "Heading" in style_name:
                level = style_name.replace("Heading ", "")
                try:
                    level_num = int(level)
                    header_prefix = "#" * min(level_num + 1, 6)
                    content_parts.append(f"\n{header_prefix} {text}")
                    current_section = text
                except:
                    content_parts.append(f"\n## {text}")
            elif style_name == "Title":
                content_parts.append(f"\n# {text}")
            else:
                # Regular paragraphs
                if current_section:
                    content_parts.append(f"\n{text}")
                else:
                    content_parts.append(text)

        # Extract tables with better formatting
        for i, table in enumerate(doc.tables):
            content_parts.append(f"\n## Table {i + 1}")
            
            # Extract headers if first row looks like headers
            if table.rows:
                first_row = table.rows[0]
                headers = [cell.text.strip() for cell in first_row.cells]
                
                # Check if this looks like headers (short text, different formatting)
                is_header_row = all(len(h) < 50 and h for h in headers)
                
                if is_header_row:
                    content_parts.append("| " + " | ".join(headers) + " |")
                    content_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
                    start_row = 1
                else:
                    start_row = 0
                
                # Extract data rows
                for row in table.rows[start_row:]:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_cells):  # Skip empty rows
                        content_parts.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(content_parts)

    except ImportError:
        return f"# Word Document: {file_name}\n\nError: python-docx not installed. Install with: pip install python-docx"
    except Exception as e:
        return f"# Word Document: {file_name}\n\nError extracting Word content: {str(e)}"

def format_csv_content(content, file_name):
    """Format CSV content with analysis"""
    try:
        import pandas as pd
        from io import StringIO

        df = pd.read_csv(StringIO(content))

        formatted = f"# CSV File: {file_name}\n\n"
        formatted += f"## Dataset Information\n"
        formatted += f"- Rows: {len(df)}\n"
        formatted += f"- Columns: {len(df.columns)}\n"
        formatted += f"- Column Names: {', '.join(df.columns.tolist())}\n\n"

        # Add sample data
        formatted += f"## Sample Data (First 5 Rows)\n"
        formatted += df.head().to_string() + "\n\n"

        # Add data types
        formatted += f"## Data Types\n"
        for col, dtype in df.dtypes.items():
            formatted += f"- {col}: {dtype}\n"

        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            formatted += f"\n## Numeric Summary\n"
            formatted += df[numeric_cols].describe().to_string()

        formatted += f"\n\n## Full Data\n```csv\n{content}\n```"

        return formatted

    except Exception as e:
        return f"# CSV File: {file_name}\n\nError processing CSV: {str(e)}\n\n## Raw Content\n```csv\n{content}\n```"

def format_python_content(content, file_name):
    """Format Python code with documentation extraction"""
    formatted = f"# Python File: {file_name}\n\n"

    # Extract docstrings and comments
    lines = content.split('\n')
    in_docstring = False
    docstring_content = []

    for line in lines:
        stripped = line.strip()
        if '"""' in stripped or "'''" in stripped:
            if in_docstring:
                in_docstring = False
                if docstring_content:
                    formatted += "## Documentation:\n" + '\n'.join(docstring_content) + "\n\n"
                    docstring_content = []
            else:
                in_docstring = True
        elif in_docstring:
            docstring_content.append(stripped)
        elif stripped.startswith('#'):
            formatted += f"## Comment: {stripped[1:].strip()}\n"

    formatted += f"\n## Full Code:\n```python\n{content}\n```"
    return formatted

def format_javascript_content(content, file_name):
    """Format JavaScript/TypeScript code with documentation extraction"""
    formatted = f"# JavaScript/TypeScript File: {file_name}\n\n"

    # Extract JSDoc comments and regular comments
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('/**') or stripped.startswith('*') or stripped.startswith('//'):
            formatted += f"## Documentation: {stripped}\n"

    formatted += f"\n## Full Code:\n```javascript\n{content}\n```"
    return formatted

def format_java_content(content, file_name):
    """Format Java code with documentation extraction"""
    formatted = f"# Java File: {file_name}\n\n"

    # Extract class and method information
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if 'class ' in stripped and 'public' in stripped:
            formatted += f"## Class: {stripped}\n"
        elif 'public ' in stripped and '(' in stripped:
            formatted += f"## Method: {stripped}\n"
        elif stripped.startswith('/**') or stripped.startswith('*') or stripped.startswith('//'):
            formatted += f"## Documentation: {stripped}\n"

    formatted += f"\n## Full Code:\n```java\n{content}\n```"
    return formatted

def format_cpp_content(content, file_name):
    """Format C/C++ code with documentation extraction"""
    formatted = f"# C/C++ File: {file_name}\n\n"

    # Extract function declarations and comments
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('//') or stripped.startswith('/*'):
            formatted += f"## Documentation: {stripped}\n"
        elif '(' in stripped and ('{' in stripped or ';' in stripped):
            formatted += f"## Function: {stripped}\n"

    formatted += f"\n## Full Code:\n```cpp\n{content}\n```"
    return formatted

def format_go_content(content, file_name):
    """Format Go code with documentation extraction"""
    formatted = f"# Go File: {file_name}\n\n"

    # Extract package, function, and comment information
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('package '):
            formatted += f"## Package: {stripped}\n"
        elif stripped.startswith('func '):
            formatted += f"## Function: {stripped}\n"
        elif stripped.startswith('//'):
            formatted += f"## Documentation: {stripped}\n"

    formatted += f"\n## Full Code:\n```go\n{content}\n```"
    return formatted

def format_rust_content(content, file_name):
    """Format Rust code with documentation extraction"""
    formatted = f"# Rust File: {file_name}\n\n"

    # Extract function, struct, and comment information
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('fn '):
            formatted += f"## Function: {stripped}\n"
        elif stripped.startswith('struct '):
            formatted += f"## Struct: {stripped}\n"
        elif stripped.startswith('//') or stripped.startswith('///'):
            formatted += f"## Documentation: {stripped}\n"

    formatted += f"\n## Full Code:\n```rust\n{content}\n```"
    return formatted

def format_markdown_content(content, file_name):
    """Format Markdown content with structure preservation"""
    return f"# Markdown Document: {file_name}\n\n{content}"

def format_config_content(content, file_name, file_extension):
    """Format configuration files"""
    formatted = f"# Configuration File: {file_name}\n"
    formatted += f"Type: {file_extension.upper()}\n\n"
    formatted += f"## Configuration Content:\n```{file_extension[1:]}\n{content}\n```"
    return formatted

def chunk_text(text, chunk_size=1000, chunk_overlap=200, chunk_strategy='language-specific', file_extension=None):
    """
    Split text into chunks using different strategies optimized for RAG systems
    
    This is a critical function for the RAG (Retrieval-Augmented Generation) pipeline
    that breaks down large text documents into manageable chunks for vector indexing
    and semantic search. Different chunking strategies preserve different aspects
    of the document structure and meaning.
    
    Chunking Strategies:
        - language-specific: Uses syntax-aware chunking for code files (functions, classes)
        - semantic: Groups related sentences and paragraphs together  
        - size-based: Simple character-count based chunking with overlap
        
    Args:
        text (str): The text content to chunk
        chunk_size (int, optional): Target size for each chunk in characters. Defaults to 1000.
        chunk_overlap (int, optional): Number of overlapping characters between chunks. Defaults to 200.
        chunk_strategy (str, optional): Strategy to use ('language-specific', 'semantic', 'size'). Defaults to 'language-specific'.
        file_extension (str, optional): File extension for language-specific chunking (e.g., '.py', '.js'). Defaults to None.
        
    Returns:
        list[str]: List of text chunks with appropriate overlap
        
    Raises:
        ValueError: If chunk_size is too small or chunk_overlap >= chunk_size
        
    Example:
        >>> text = "def function1():\n    pass\n\ndef function2():\n    return True"
        >>> chunks = chunk_text(text, chunk_size=500, chunk_strategy='language-specific', file_extension='.py')
        >>> print(f"Created {len(chunks)} chunks")
        
    Note:
        - Overlap helps maintain context between chunks for better retrieval
        - Language-specific chunking preserves code structure and function boundaries
        - Semantic chunking groups related content for better coherence
        - For small texts (< chunk_size), returns single chunk without processing
    """
    app_logger.debug(f"Starting text chunking: {len(text)} characters, strategy='{chunk_strategy}', size={chunk_size}, overlap={chunk_overlap}")
    
    # Validate input parameters
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be less than chunk_size")
    
    # Return single chunk for small texts
    if len(text) <= chunk_size:
        app_logger.debug("Text is smaller than chunk_size, returning single chunk")
        return [text]

    # Apply chunking strategy
    try:
        if chunk_strategy == 'language_specific' and file_extension:
            app_logger.debug(f"Using language-specific chunking for {file_extension}")
            chunks = chunk_by_language(text, file_extension, chunk_size, chunk_overlap)
        elif chunk_strategy == 'semantic':
            app_logger.debug("Using semantic chunking")
            chunks = chunk_semantically(text, chunk_size, chunk_overlap)
        else:
            app_logger.debug("Using size-based chunking")
            chunks = chunk_by_size(text, chunk_size, chunk_overlap)
            
        app_logger.info(f"Successfully created {len(chunks)} chunks using {chunk_strategy} strategy")
        
        # Log chunk statistics
        if chunks:
            avg_chunk_size = sum(len(chunk) for chunk in chunks) / len(chunks)
            min_chunk_size = min(len(chunk) for chunk in chunks)
            max_chunk_size = max(len(chunk) for chunk in chunks)
            app_logger.debug(f"Chunk stats - Average: {avg_chunk_size:.1f}, Min: {min_chunk_size}, Max: {max_chunk_size}")
        
        return chunks
        
    except Exception as e:
        app_logger.error(f"Error during text chunking: {e}")
        log_error_with_context(e, {
            "text_length": len(text),
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "chunk_strategy": chunk_strategy,
            "file_extension": file_extension
        })
        # Fallback to simple size-based chunking
        app_logger.warning("Falling back to size-based chunking due to error")
        return chunk_by_size(text, chunk_size, chunk_overlap)

def chunk_by_language(text, file_extension, chunk_size=1000, chunk_overlap=200):
    """Language-specific chunking for code files"""
    if file_extension in ['.py']:
        return chunk_python_code(text, chunk_size, chunk_overlap)
    elif file_extension in ['.js', '.ts']:
        return chunk_javascript_code(text, chunk_size, chunk_overlap)
    elif file_extension in ['.java']:
        return chunk_java_code(text, chunk_size, chunk_overlap)
    elif file_extension in ['.cpp', '.c', '.h', '.hpp']:
        return chunk_cpp_code(text, chunk_size, chunk_overlap)
    elif file_extension in ['.go']:
        return chunk_go_code(text, chunk_size, chunk_overlap)
    elif file_extension in ['.rs']:
        return chunk_rust_code(text, chunk_size, chunk_overlap)
    elif file_extension == '.md':
        return chunk_markdown(text, chunk_size, chunk_overlap)
    else:
        return chunk_semantically(text, chunk_size, chunk_overlap)

def chunk_python_code(text, chunk_size, chunk_overlap):
    """Chunk Python code by functions, classes, and logical blocks"""
    chunks = []
    lines = text.split('\n')
    current_chunk = []
    current_size = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        # Check for function or class definitions
        if line.strip().startswith(('def ', 'class ', 'async def ')):
            # If we have content, save current chunk
            if current_chunk and current_size > chunk_size // 2:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_size = 0

            # Add the function/class definition and its body
            current_chunk.append(line)
            current_size += len(line)
            i += 1

            # Get the indentation level
            indent_level = len(line) - len(line.lstrip())

            # Add all lines that belong to this function/class
            while i < len(lines):
                next_line = lines[i]
                if next_line.strip() == '':
                    current_chunk.append(next_line)
                    current_size += len(next_line)
                    i += 1
                elif len(next_line) - len(next_line.lstrip()) > indent_level:
                    current_chunk.append(next_line)
                    current_size += len(next_line)
                    i += 1
                else:
                    break
        else:
            current_chunk.append(line)
            current_size += len(line)
            i += 1

            # If chunk is getting too large, split it
            if current_size > chunk_size:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_size = 0

    # Add remaining content
    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks if chunks else [text]

def chunk_javascript_code(text, chunk_size, chunk_overlap):
    """Chunk JavaScript/TypeScript code by functions and classes"""
    chunks = []
    lines = text.split('\n')
    current_chunk = []
    current_size = 0
    brace_count = 0

    for line in lines:
        current_chunk.append(line)
        current_size += len(line)

        # Count braces to track function/class boundaries
        brace_count += line.count('{') - line.count('}')

        # If we're at a function/class boundary and chunk is large enough
        if brace_count == 0 and current_size > chunk_size // 2:
            chunks.append('\n'.join(current_chunk))
            current_chunk = []
            current_size = 0
        elif current_size > chunk_size:
            chunks.append('\n'.join(current_chunk))
            current_chunk = []
            current_size = 0
            brace_count = 0

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks if chunks else [text]

def chunk_java_code(text, chunk_size, chunk_overlap):
    """Chunk Java code by methods and classes"""
    return chunk_javascript_code(text, chunk_size, chunk_overlap)  # Similar structure

def chunk_cpp_code(text, chunk_size, chunk_overlap):
    """Chunk C/C++ code by functions"""
    return chunk_javascript_code(text, chunk_size, chunk_overlap)  # Similar structure

def chunk_go_code(text, chunk_size, chunk_overlap):
    """Chunk Go code by functions and structs"""
    return chunk_javascript_code(text, chunk_size, chunk_overlap)  # Similar structure

def chunk_rust_code(text, chunk_size, chunk_overlap):
    """Chunk Rust code by functions and structs"""
    return chunk_javascript_code(text, chunk_size, chunk_overlap)  # Similar structure

def chunk_markdown(text, chunk_size, chunk_overlap):
    """Chunk Markdown by headers and sections"""
    chunks = []
    lines = text.split('\n')
    current_chunk = []
    current_size = 0

    for line in lines:
        # If we hit a header and current chunk is substantial
        if line.startswith('#') and current_chunk and current_size > chunk_size // 2:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_size = len(line)
        else:
            current_chunk.append(line)
            current_size += len(line)

            if current_size > chunk_size:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_size = 0

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks if chunks else [text]

def chunk_semantically(text, chunk_size, chunk_overlap):
    """Semantic chunking by paragraphs and sentences"""
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at paragraph boundaries first
        if end < len(text):
            paragraph_end = text.rfind('\n\n', start, end)
            if paragraph_end > start + chunk_size // 2:
                end = paragraph_end + 2
            else:
                # Try sentence boundaries
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size - 100:
                    end = sentence_end + 1
                else:
                    # Fall back to word boundaries
                    word_end = text.rfind(' ', start, end)
                    if word_end > start + chunk_size - 50:
                        end = word_end

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - chunk_overlap
        if start >= len(text):
            break

    return chunks

def chunk_by_size(text, chunk_size, chunk_overlap):
    """Simple size-based chunking"""
    return chunk_semantically(text, chunk_size, chunk_overlap)

def generate_intelligent_response(query, relevant_chunks, all_chunks):
    """Generate intelligent responses based on document content"""
    query_lower = query.lower()

    # Check for summary/overview requests
    summary_keywords = {'summary', 'summarize', 'overview', 'main points', 'key points', 'what is', 'about', 'describe'}
    is_summary_request = any(keyword in query_lower for keyword in summary_keywords)

    if not relevant_chunks and not all_chunks:
        return "I don't have any documents to analyze. Please upload some documents to your context first."

    if is_summary_request and all_chunks:
        return generate_document_summary(all_chunks)
    elif relevant_chunks:
        return generate_contextual_response(query, relevant_chunks)
    else:
        return f"I couldn't find specific information about '{query}' in your uploaded documents. Try rephrasing your question or asking for a general summary."

def generate_document_summary(all_chunks):
    """Generate a comprehensive summary of all documents"""
    if not all_chunks:
        return "No documents found to summarize."

    # Group chunks by file
    files_content = {}
    for chunk in all_chunks:
        file_name = chunk['file_name']
        if file_name not in files_content:
            files_content[file_name] = []
        files_content[file_name].append(chunk['content'])

    summary_parts = ["📄 **Document Summary**\n"]

    for file_name, chunks in files_content.items():
        # Combine all chunks for this file
        full_content = ' '.join(chunks)

        # Extract key information
        key_points = extract_key_points(full_content)

        summary_parts.append(f"\n**📁 {file_name}:**")
        if key_points:
            for point in key_points[:3]:  # Top 3 key points per file
                summary_parts.append(f"• {point}")
        else:
            # Fallback: use first chunk as summary
            first_chunk = chunks[0][:300] + "..." if len(chunks[0]) > 300 else chunks[0]
            summary_parts.append(f"• {first_chunk}")

    # Add overall statistics
    total_files = len(files_content)
    total_chunks = len(all_chunks)

    summary_parts.append(f"\n📊 **Statistics:**")
    summary_parts.append(f"• Total files: {total_files}")
    summary_parts.append(f"• Total sections: {total_chunks}")

    return '\n'.join(summary_parts)

def generate_contextual_response(query, relevant_chunks):
    """Generate response based on relevant chunks"""
    if not relevant_chunks:
        return f"I couldn't find specific information about '{query}' in your documents."

    response_parts = [f"📋 **Based on your documents, here's what I found about '{query}':**\n"]

    # Group by file for better organization
    files_info = {}
    for chunk in relevant_chunks[:5]:  # Top 5 most relevant chunks
        file_name = chunk['file_name']
        if file_name not in files_info:
            files_info[file_name] = []
        files_info[file_name].append(chunk['content'])

    for file_name, contents in files_info.items():
        response_parts.append(f"\n**📁 From {file_name}:**")
        for content in contents:
            # Truncate long content but keep it meaningful
            if len(content) > 400:
                content = content[:400] + "..."
            response_parts.append(f"• {content}")

    return '\n'.join(response_parts)

def extract_key_points(text):
    """Extract key points from text using simple heuristics"""
    if not text:
        return []

    # Split into sentences
    sentences = re.split(r'[.!?]+', text)
    key_points = []

    # Look for sentences that might be key points
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 20 or len(sentence) > 200:
            continue

        # Prioritize sentences with key indicators
        key_indicators = ['purpose', 'goal', 'objective', 'main', 'key', 'important', 'summary', 'overview']
        if any(indicator in sentence.lower() for indicator in key_indicators):
            key_points.append(sentence)
        elif len(key_points) < 3:  # Add first few sentences as fallback
            key_points.append(sentence)

    return key_points[:5]  # Return top 5 key points

def generate_vector_search_response(query, context_ids):
    """Generate response using vector search for repository contexts"""
    try:
        from services.gemini_service import GeminiService
        import faiss
        import numpy as np

        # Initialize services
        gemini_service = GeminiService()

        # Create query embedding
        query_embedding = gemini_service.create_query_embedding(query)
        query_vector = np.array([query_embedding], dtype=np.float32)
        faiss.normalize_L2(query_vector)

        # Search across all context vector stores
        all_results = []

        for context_id in context_ids:
            try:
                # Load FAISS index
                vector_store_dir = os.path.join('vector_stores', str(context_id))
                index_path = os.path.join(vector_store_dir, 'index.faiss')
                metadata_path = os.path.join(vector_store_dir, 'metadata.json')

                if not os.path.exists(index_path) or not os.path.exists(metadata_path):
                    continue

                # Load index and metadata
                index = faiss.read_index(index_path)
                with open(metadata_path, 'r') as f:
                    chunks_metadata = json.load(f)

                # Search for similar chunks
                k = min(5, index.ntotal)  # Top 5 or all available
                scores, indices = index.search(query_vector, k)

                # Collect results
                for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
                    if idx < len(chunks_metadata) and score > 0.3:  # Similarity threshold
                        chunk_data = chunks_metadata[idx]
                        all_results.append({
                            'content': chunk_data['content'],
                            'source': chunk_data['file_name'],
                            'score': float(score),
                            'context_id': context_id,
                            'metadata': chunk_data['metadata']
                        })

            except Exception as e:
                print(f"Error searching context {context_id}: {e}")
                continue

        # Sort by relevance score
        all_results.sort(key=lambda x: x['score'], reverse=True)
        top_results = all_results[:10]  # Top 10 results

        if not top_results:
            return "I couldn't find relevant information in the repository code. Please try rephrasing your question or check if the repository has been processed correctly."

        # Generate response using Gemini
        response_data = gemini_service.generate_response(
            query=query,
            context_chunks=top_results,
            max_tokens=2048
        )

        # Format response with code context
        formatted_response = f"🔍 **Based on the repository code analysis:**\n\n{response_data['content']}\n\n"

        # Add source references
        if top_results:
            formatted_response += "📁 **Sources:**\n"
            seen_sources = set()
            for result in top_results[:5]:
                source = result['source']
                if source not in seen_sources:
                    language = result['metadata'].get('language', 'unknown')
                    formatted_response += f"• `{source}` ({language})\n"
                    seen_sources.add(source)

        return formatted_response

    except Exception as e:
        print(f"Error in vector search response: {e}")
        return f"I encountered an error while searching the repository: {str(e)}. Please try again or check if the repository has been processed correctly."

def process_context_files(context_id, file_paths, chunk_strategy='language-specific'):
    """Process uploaded files and create chunks with language-specific handling"""
    all_chunks = []

    for file_path in file_paths:
        try:
            # Extract text from file with language-specific formatting
            text = extract_text_from_file(file_path)
            file_extension = Path(file_path).suffix.lower()
            file_name = Path(file_path).name

            # Determine chunk size based on file type
            if file_extension in ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.go', '.rs']:
                chunk_size = 1500  # Larger chunks for code files
                chunk_overlap = 300
            else:
                chunk_size = 1000  # Standard size for text files
                chunk_overlap = 200

            # Create chunks using specified strategy
            chunks = chunk_text(text, chunk_size, chunk_overlap, chunk_strategy, file_extension)

            # Add metadata to chunks
            for i, chunk in enumerate(chunks):
                # Analyze chunk content for additional metadata
                chunk_metadata = analyze_chunk_content(chunk, file_extension)

                chunk_data = {
                    'context_id': context_id,
                    'file_name': file_name,
                    'chunk_index': i,
                    'content': chunk,
                    'metadata': {
                        'file_path': file_path,
                        'file_type': file_extension,
                        'chunk_size': len(chunk),
                        'chunk_strategy': chunk_strategy,
                        'language': detect_language(file_extension),
                        'content_type': chunk_metadata['content_type'],
                        'has_code': chunk_metadata['has_code'],
                        'has_documentation': chunk_metadata['has_documentation'],
                        'complexity_score': chunk_metadata['complexity_score']
                    }
                }
                all_chunks.append(chunk_data)

        except Exception as e:
            print(f"Error processing file {file_path}: {str(e)}")
            continue

    return all_chunks

def analyze_chunk_content(chunk, file_extension):
    """Analyze chunk content to extract metadata"""
    metadata = {
        'content_type': 'text',
        'has_code': False,
        'has_documentation': False,
        'complexity_score': 0
    }

    chunk_lower = chunk.lower()

    # Detect content type
    if file_extension in ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.go', '.rs']:
        metadata['content_type'] = 'code'
        metadata['has_code'] = True

        # Check for documentation
        if any(marker in chunk for marker in ['"""', "'''", '/**', '///', '##']):
            metadata['has_documentation'] = True

        # Simple complexity scoring
        complexity_indicators = ['if ', 'for ', 'while ', 'try ', 'catch ', 'function ', 'class ', 'def ']
        metadata['complexity_score'] = sum(chunk_lower.count(indicator) for indicator in complexity_indicators)

    elif file_extension == '.md':
        metadata['content_type'] = 'documentation'
        metadata['has_documentation'] = True

        # Count headers and lists for complexity
        metadata['complexity_score'] = chunk.count('#') + chunk.count('- ') + chunk.count('* ')

    elif file_extension in ['.json', '.yaml', '.yml']:
        metadata['content_type'] = 'configuration'
        metadata['complexity_score'] = chunk.count(':') + chunk.count('-')

    return metadata

def detect_language(file_extension):
    """Detect programming language from file extension"""
    language_map = {
        '.py': 'python',
        '.js': 'javascript',
        '.ts': 'typescript',
        '.java': 'java',
        '.cpp': 'cpp',
        '.c': 'c',
        '.h': 'c',
        '.hpp': 'cpp',
        '.go': 'go',
        '.rs': 'rust',
        '.rb': 'ruby',
        '.php': 'php',
        '.cs': 'csharp',
        '.kt': 'kotlin',
        '.swift': 'swift',
        '.md': 'markdown',
        '.txt': 'text',
        '.json': 'json',
        '.yaml': 'yaml',
        '.yml': 'yaml'
    }

    return language_map.get(file_extension, 'unknown')

def process_repository_context(context_id, config):
    """Process repository context by cloning and analyzing code"""
    try:
        from services.repository_service import RepositoryService
        from services.gemini_service import GeminiService
        import faiss
        import numpy as np

        # Get context
        context = Context.query.get(context_id)
        if not context:
            print(f"Context {context_id} not found")
            return

        # Update status
        context.status = 'processing'
        context.progress = 10
        db.session.commit()

        # Get repository configuration - check multiple possible keys
        repo_url = config.get('repo_url') or config.get('url') or config.get('repository_url')
        branch = config.get('branch', 'main')

        print(f"Repository config received: {config}")
        print(f"Extracted repo_url: {repo_url}")
        print(f"Branch: {branch}")

        if not repo_url:
            context.status = 'error'
            context.error_message = f'Repository URL not provided. Config received: {config}'
            db.session.commit()
            return

        print(f"Processing repository: {repo_url}")

        # Clone repository
        repo_service = RepositoryService()
        clone_result = repo_service.clone_repository_direct(repo_url, branch)

        if not clone_result['success']:
            context.status = 'error'
            context.error_message = f"Failed to clone repository: {clone_result['error']}"
            db.session.commit()
            return

        context.progress = 30
        db.session.commit()

        # Get repository info and processable files
        repo_info = clone_result['repo_info']
        clone_path = clone_result['clone_path']
        processable_files = repo_info['file_analysis']['processable_files']

        print(f"Found {len(processable_files)} processable files")

        # Process files and create chunks
        chunk_strategy = config.get('chunk_strategy', 'language_specific')
        all_chunks = []

        for i, file_info in enumerate(processable_files[:50]):  # Limit to 50 files
            try:
                file_path = file_info['path']

                # Extract text from file
                text = extract_text_from_file(file_path)
                file_extension = Path(file_path).suffix.lower()

                # Create chunks
                chunks = chunk_text(text, 1500, 300, chunk_strategy, file_extension)

                # Add metadata to chunks
                for j, chunk in enumerate(chunks):
                    chunk_data = {
                        'context_id': context_id,
                        'file_name': file_info['relative_path'],
                        'chunk_index': j,
                        'content': chunk,
                        'metadata': {
                            'file_path': file_info['relative_path'],
                            'file_type': file_extension,
                            'language': file_info['language'],
                            'chunk_size': len(chunk),
                            'chunk_strategy': chunk_strategy,
                            'repository_url': repo_url,
                            'repository_branch': branch,
                            'project_type': repo_info['project_info']['primary_type']
                        }
                    }
                    all_chunks.append(chunk_data)

            except Exception as e:
                print(f"Error processing file {file_info['path']}: {e}")
                continue

            # Update progress
            progress = 30 + int((i / len(processable_files)) * 40)
            context.progress = progress
            db.session.commit()

        print(f"Created {len(all_chunks)} chunks")

        # Create embeddings and store in vector database
        if all_chunks:
            context.progress = 70
            db.session.commit()

            # Initialize Gemini service for embeddings
            gemini_service = GeminiService()

            # Create embeddings for chunks in batches
            chunk_texts = [chunk['content'] for chunk in all_chunks]
            print(f"Creating embeddings for {len(chunk_texts)} chunks...")

            # Process in smaller batches to avoid timeouts
            batch_size = 5
            embeddings = []

            for i in range(0, len(chunk_texts), batch_size):
                batch = chunk_texts[i:i + batch_size]
                print(f"Processing batch {i//batch_size + 1}/{(len(chunk_texts) + batch_size - 1)//batch_size}")

                batch_embeddings = gemini_service.create_embeddings(batch, "retrieval_document")
                embeddings.extend(batch_embeddings)

                # Update progress
                progress = 70 + int((i / len(chunk_texts)) * 15)
                context.progress = progress
                db.session.commit()

            # Create FAISS index
            embedding_dim = len(embeddings[0]) if embeddings else 768
            index = faiss.IndexFlatIP(embedding_dim)  # Inner product for cosine similarity

            # Normalize embeddings and add to index
            embeddings_array = np.array(embeddings, dtype=np.float32)
            faiss.normalize_L2(embeddings_array)
            index.add(embeddings_array)

            # Save FAISS index
            vector_store_dir = os.path.join('vector_stores', str(context_id))
            os.makedirs(vector_store_dir, exist_ok=True)
            vector_store_path = os.path.join(vector_store_dir, 'index.faiss')
            faiss.write_index(index, vector_store_path)

            # Save chunk metadata
            metadata_path = os.path.join(vector_store_dir, 'metadata.json')
            with open(metadata_path, 'w') as f:
                json.dump(all_chunks, f, indent=2)

            context.progress = 90
            db.session.commit()

        # Store chunks in database
        for chunk_data in all_chunks:
            chunk = TextChunk(
                context_id=chunk_data['context_id'],
                file_name=chunk_data['file_name'],
                chunk_index=chunk_data['chunk_index'],
                content=chunk_data['content']
            )
            chunk.set_file_info(chunk_data['metadata'])
            db.session.add(chunk)

        # Update context with final statistics
        context.total_chunks = len(all_chunks)
        context.total_tokens = sum(len(chunk['content'].split()) for chunk in all_chunks)
        context.status = 'ready'
        context.progress = 100
        db.session.commit()

        # Cleanup cloned repository
        repo_service.cleanup_repository(clone_path)

        print(f"Repository processing completed for context {context_id}")

    except Exception as e:
        print(f"Error processing repository context {context_id}: {e}")
        context = Context.query.get(context_id)
        if context:
            context.status = 'error'
            context.error_message = str(e)
            db.session.commit()

# User model is imported from models.py

# Simple Context model
# Context model is imported from models.py - no need to redefine here

# ChatSession model is imported from models.py

# Message model is imported from models.py
# TextChunk model is imported from models.py

# Routes

@app.route('/api/auth/register', methods=['POST'])
def register():
    try:
        data = request.get_json()
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')

        if not username or not email or not password:
            return jsonify({'error': 'Missing required fields'}), 400

        # Check if user exists
        if User.query.filter_by(username=username).first():
            return jsonify({'error': 'Username already exists'}), 400
        
        if User.query.filter_by(email=email).first():
            return jsonify({'error': 'Email already exists'}), 400

        # Create user
        user = User(username=username, email=email)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()

        # Create access token (convert user.id to string)
        access_token = create_access_token(identity=str(user.id))

        return jsonify({
            'access_token': access_token,
            'user': user.to_dict()
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """
    Authenticate user and generate JWT access token
    
    This endpoint handles user authentication by validating credentials
    and generating a JWT token for API access. The token includes the
    user ID and has a 24-hour expiration time.
    
    Request Body:
        {
            "username": "string (required) - User's username or email",
            "password": "string (required) - User's password"
        }
        
    Returns:
        200: Authentication successful
        {
            "access_token": "JWT token for API access",
            "user": {
                "id": "User ID",
                "username": "Username", 
                "email": "Email address",
                "is_admin": "Boolean admin status"
            }
        }
        
        400: Missing required fields
        401: Invalid credentials
        500: Internal server error
        
    Security:
        - Passwords are hashed using Werkzeug security
        - JWT tokens expire after 24 hours
        - Failed login attempts are logged for security monitoring
        - No sensitive information is exposed in error messages
        
    Example:
        POST /api/auth/login
        {
            "username": "admin",
            "password": "admin123"
        }
    """
    client_ip = request.remote_addr
    user_agent = request.headers.get('User-Agent', 'Unknown')
    app_logger.info(f"Login attempt from IP: {client_ip}, User-Agent: {user_agent}")
    
    try:
        # Parse request data
        data = request.get_json()
        if not data:
            app_logger.warning(f"Login attempt with no JSON data from {client_ip}")
            return jsonify({'error': 'Invalid JSON data'}), 400
            
        username = data.get('username')
        password = data.get('password')

        # Validate required fields
        if not username or not password:
            app_logger.warning(f"Login attempt with missing credentials from {client_ip}")
            return jsonify({'error': 'Missing username or password'}), 400

        # Sanitize username for logging (avoid logging sensitive data)
        username_safe = username[:3] + "*" * (len(username) - 3) if len(username) > 3 else "***"
        app_logger.debug(f"Processing login for user: {username_safe}")

        # Find user by username
        user = User.query.filter_by(username=username).first()
        if not user:
            app_logger.warning(f"Login failed: User not found for username {username_safe} from {client_ip}")
            return jsonify({'error': 'Invalid credentials'}), 401

        # Verify password
        if not user.check_password(password):
            app_logger.warning(f"Login failed: Invalid password for user {username_safe} from {client_ip}")
            return jsonify({'error': 'Invalid credentials'}), 401

        # Generate access token
        access_token = create_access_token(identity=str(user.id))
        
        # Log successful authentication
        app_logger.info(f"Successful login for user {user.username} (ID: {user.id}) from {client_ip}")

        return jsonify({
            'access_token': access_token,
            'user': user.to_dict()
        }), 200

    except Exception as e:
        app_logger.error(f"Error during login process from {client_ip}: {str(e)}")
        log_error_with_context(e, {
            "endpoint": "/api/auth/login",
            "client_ip": client_ip,
            "user_agent": user_agent,
            "username": username_safe if 'username' in locals() else "unknown"
        })
        return jsonify({'error': 'Authentication failed'}), 500

@app.route('/api/auth/profile', methods=['GET'])
@jwt_required()
def get_profile():
    try:
        user_id = get_current_user_id()
        user = db.session.get(User, user_id)

        if not user:
            return jsonify({'error': 'User not found'}), 404

        return jsonify({'user': user.to_dict()}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts', methods=['GET'])
@jwt_required()
def get_contexts():
    try:
        user_id = get_current_user_id()
        contexts = Context.query.filter_by(user_id=user_id).all()
        return jsonify({'contexts': [ctx.to_dict() for ctx in contexts]}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts', methods=['POST'])
@jwt_required()
def create_context():
    try:
        user_id = get_current_user_id()
        data = request.get_json()

        # Extract configuration values
        chunk_strategy = data.get('chunk_strategy', 'language-specific')
        embedding_model = data.get('embedding_model', 'text-embedding-004')

        context = Context(
            name=data.get('name'),
            description=data.get('description', ''),
            source_type=data.get('source_type'),
            chunk_strategy=chunk_strategy,
            embedding_model=embedding_model,
            user_id=user_id,
            status='pending'  # Start as pending for processing
        )

        # Set the full configuration including source-specific config
        full_config = data.get('config', {})
        full_config.update({
            'chunk_strategy': chunk_strategy,
            'embedding_model': embedding_model
        })
        context.set_config(full_config)

        db.session.add(context)
        db.session.commit()

        # Process repository if source type is repo
        if data.get('source_type') == 'repo':
            # Extract repository config from the data
            repo_config = data.get('repo_config', {})
            if not repo_config and data.get('config'):
                # Fallback to config if repo_config is not present
                repo_config = data.get('config', {})
            process_repository_context(context.id, repo_config)

        return jsonify({'context': context.to_dict()}), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts/<int:context_id>', methods=['GET'])
@jwt_required()
def get_context(context_id):
    try:
        user_id = get_current_user_id()
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()

        if not context:
            return jsonify({'error': 'Context not found'}), 404

        # Get chunks for this context
        chunks = TextChunk.query.filter_by(context_id=context_id).all()

        # Group chunks by file to create documents data
        documents = {}
        for chunk in chunks:
            file_name = chunk.file_name
            if file_name not in documents:
                documents[file_name] = {
                    'id': f"{context_id}_{file_name}",
                    'filename': file_name,
                    'file_type': chunk.get_file_info().get('file_type', 'unknown'),
                    'file_size': chunk.get_file_info().get('chunk_size', 0),
                    'chunks_count': 0,
                    'tokens_count': 0,
                    'language': None,
                    'processed_at': chunk.created_at.isoformat()
                }
            documents[file_name]['chunks_count'] += 1
            documents[file_name]['tokens_count'] += len(chunk.content.split())

        context_data = context.to_dict()
        context_data['documents'] = list(documents.values())
        context_data['chunks'] = [chunk.to_dict() for chunk in chunks]

        return jsonify({'context': context_data}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts/<int:context_id>', methods=['PUT'])
@jwt_required()
def update_context(context_id):
    try:
        user_id = get_current_user_id()
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()

        if not context:
            return jsonify({'error': 'Context not found'}), 404

        data = request.get_json()
        if 'name' in data:
            context.name = data['name']
        if 'description' in data:
            context.description = data['description']

        db.session.commit()
        return jsonify({'context': context.to_dict()}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts/<int:context_id>', methods=['DELETE'])
@jwt_required()
def delete_context(context_id):
    try:
        user_id = get_current_user_id()

        # Get context and verify ownership
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()
        if not context:
            return jsonify({'error': 'Context not found'}), 404

        cleanup_stats = {
            'vector_stores_deleted': 0,
            'documents_deleted': 0,
            'files_deleted': 0,
            'chunks_deleted': 0
        }

        try:
            # Delete vector store files if they exist
            vector_store_dir = os.path.join('vector_stores', str(context_id))
            if os.path.exists(vector_store_dir):
                import shutil
                shutil.rmtree(vector_store_dir)
                cleanup_stats['vector_stores_deleted'] = 1
        except Exception as e:
            print(f"Error deleting vector store: {e}")

        try:
            # Delete uploaded files if they exist
            upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], str(context_id))
            if os.path.exists(upload_dir):
                import shutil
                shutil.rmtree(upload_dir)
                cleanup_stats['files_deleted'] = 1
        except Exception as e:
            print(f"Error deleting upload files: {e}")

        # Delete text chunks if they exist
        try:
            chunks_deleted = TextChunk.query.filter_by(context_id=context_id).delete()
            cleanup_stats['chunks_deleted'] = chunks_deleted
        except Exception as e:
            print(f"Error deleting chunks: {e}")

        # Delete the context (this will cascade delete documents)
        db.session.delete(context)
        db.session.commit()

        return jsonify({
            'message': f'Context "{context.name}" deleted successfully',
            'cleanup_stats': cleanup_stats
        }), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error during context cleanup: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts/<int:context_id>/reprocess', methods=['POST'])
@jwt_required()
def reprocess_context(context_id):
    try:
        user_id = get_current_user_id()
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()

        if not context:
            return jsonify({'error': 'Context not found'}), 404

        # Update context status to processing
        context.status = 'processing'
        context.updated_at = datetime.now(timezone.utc)
        db.session.commit()

        # Here you would typically trigger the reprocessing job
        # For now, we'll just simulate it by updating the status
        # In a real implementation, you'd use a task queue like Celery

        return jsonify({
            'message': 'Context reprocessing started',
            'context': context.to_dict()
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts/<int:context_id>/status', methods=['GET'])
@jwt_required()
def get_context_status(context_id):
    try:
        user_id = get_current_user_id()
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()

        if not context:
            return jsonify({'error': 'Context not found'}), 404

        return jsonify({
            'status': context.status,
            'progress': context.progress,
            'error_message': context.error_message,
            'total_chunks': context.total_chunks,
            'total_tokens': context.total_tokens
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contexts/<int:context_id>/chunks', methods=['GET'])
@jwt_required()
def get_context_chunks(context_id):
    try:
        user_id = get_current_user_id()
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()

        if not context:
            return jsonify({'error': 'Context not found'}), 404

        chunks = TextChunk.query.filter_by(context_id=context_id).order_by(TextChunk.file_name, TextChunk.chunk_index).all()

        return jsonify({
            'chunks': [chunk.to_dict() for chunk in chunks],
            'total_chunks': len(chunks)
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/sessions', methods=['GET'])
@jwt_required()
def get_chat_sessions():
    try:
        user_id = get_current_user_id()
        sessions = ChatSession.query.filter_by(user_id=user_id).order_by(ChatSession.updated_at.desc()).all()
        return jsonify({'sessions': [session.to_dict() for session in sessions]}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/sessions', methods=['POST'])
@jwt_required()
def create_chat_session():
    try:
        user_id = get_current_user_id()
        data = request.get_json() or {}
        
        session = ChatSession(
            title=data.get('title', f'Chat {datetime.now().strftime("%Y-%m-%d %H:%M")}'),
            user_id=user_id
        )
        
        db.session.add(session)
        db.session.commit()

        return jsonify({'session': session.to_dict()}), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/sessions/<int:session_id>', methods=['GET'])
@jwt_required()
def get_chat_session(session_id):
    try:
        user_id = get_current_user_id()
        session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()

        if not session:
            return jsonify({'error': 'Chat session not found'}), 404

        return jsonify({'session': session.to_dict()}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/sessions/<int:session_id>', methods=['PUT'])
@jwt_required()
def update_chat_session(session_id):
    try:
        user_id = get_current_user_id()
        session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()

        if not session:
            return jsonify({'error': 'Chat session not found'}), 404

        data = request.get_json()
        if 'title' in data:
            session.title = data['title']

        db.session.commit()
        return jsonify({'session': session.to_dict()}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/sessions/<int:session_id>', methods=['DELETE'])
@jwt_required()
def delete_chat_session(session_id):
    try:
        user_id = get_current_user_id()
        session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()

        if not session:
            return jsonify({'error': 'Chat session not found'}), 404

        db.session.delete(session)
        db.session.commit()
        return jsonify({'message': 'Chat session deleted successfully'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/query', methods=['POST'])
@jwt_required()
def chat_query():
    try:
        user_id = get_current_user_id()
        data = request.get_json()
        
        session_id = data.get('session_id')
        message_content = data.get('message')
        context_ids = data.get('context_ids', [])

        # Save user message
        user_message = Message(
            session_id=session_id,
            role='user',
            content=message_content
        )
        user_message.set_context_ids(context_ids)
        db.session.add(user_message)

        # Get relevant chunks from selected contexts
        relevant_chunks = []
        all_chunks = []

        if context_ids:
            for context_id in context_ids:
                # Verify user owns the context
                context = Context.query.filter_by(id=context_id, user_id=user_id).first()
                if context:
                    chunks = TextChunk.query.filter_by(context_id=context_id).all()

                    for chunk in chunks:
                        chunk_data = {
                            'content': chunk.content,
                            'file_name': chunk.file_name,
                            'context_name': context.name,
                            'chunk_index': chunk.chunk_index,
                            'relevance_score': 0
                        }
                        all_chunks.append(chunk_data)

                    # Enhanced relevance scoring
                    query_lower = message_content.lower()
                    query_words = set(query_lower.split())

                    # Check for summary/overview requests
                    summary_keywords = {'summary', 'summarize', 'overview', 'main points', 'key points', 'what is', 'about', 'describe'}
                    is_summary_request = any(keyword in query_lower for keyword in summary_keywords)

                    for chunk_data in all_chunks:
                        chunk_content_lower = chunk_data['content'].lower()

                        if is_summary_request:
                            # For summary requests, include all chunks but prioritize introductory content
                            chunk_data['relevance_score'] = 1.0
                            # Boost score for chunks that seem like introductions or overviews
                            if any(word in chunk_content_lower for word in ['introduction', 'overview', 'summary', 'abstract', 'purpose', 'goal']):
                                chunk_data['relevance_score'] = 2.0
                        else:
                            # For specific queries, use keyword matching with scoring
                            word_matches = sum(1 for word in query_words if word in chunk_content_lower)
                            if word_matches > 0:
                                chunk_data['relevance_score'] = word_matches / len(query_words)

                    # Sort by relevance and take top chunks
                    relevant_chunks = [chunk for chunk in all_chunks if chunk['relevance_score'] > 0]
                    relevant_chunks.sort(key=lambda x: x['relevance_score'], reverse=True)

        # Use vector search for repository contexts
        contexts = [db.session.get(Context, cid) for cid in context_ids]
        contexts = [ctx for ctx in contexts if ctx is not None]  # Filter out None values
        if any(ctx.source_type == 'repo' for ctx in contexts):
            ai_response = generate_vector_search_response(message_content, context_ids)
        else:
            # Generate intelligent response based on query type and relevant chunks
            ai_response = generate_intelligent_response(message_content, relevant_chunks, all_chunks)

        assistant_message = Message(
            session_id=session_id,
            role='assistant',
            content=ai_response
        )
        assistant_message.set_context_ids(context_ids)
        db.session.add(assistant_message)

        # Update session timestamp
        session = db.session.get(ChatSession, session_id)
        if session:
            session.updated_at = datetime.now(timezone.utc)

        db.session.commit()

        return jsonify({
            'message': assistant_message.to_dict(),
            'citations': []
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload/files', methods=['POST'])
@jwt_required()
def upload_files():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        if 'context_id' not in request.form:
            return jsonify({'error': 'No context_id provided'}), 400

        context_id = request.form['context_id']
        user_id = get_current_user_id()

        # Verify context belongs to user
        context = Context.query.filter_by(id=context_id, user_id=user_id).first()
        if not context:
            return jsonify({'error': 'Context not found'}), 404

        files = request.files.getlist('files')
        uploaded_files = []
        file_paths = []

        # Save files first
        for file in files:
            if file.filename == '':
                continue

            filename = secure_filename(file.filename)
            # Add timestamp to avoid conflicts
            timestamp = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_')
            filename = timestamp + filename
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            uploaded_files.append({
                'filename': file.filename,  # Original filename
                'size': os.path.getsize(file_path),
                'path': file_path
            })
            file_paths.append(file_path)

        # Process files and create chunks
        if file_paths:
            context.status = 'processing'
            db.session.commit()

            try:
                # Get chunk strategy from context configuration
                context_config = context.get_config()
                chunk_strategy = context_config.get('chunk_strategy', context.chunk_strategy or 'language-specific')

                # Process files into chunks using the specified strategy
                chunks_data = process_context_files(context_id, file_paths, chunk_strategy)

                # Save chunks to database
                for chunk_data in chunks_data:
                    chunk = TextChunk(
                        context_id=chunk_data['context_id'],
                        file_name=chunk_data['file_name'],
                        chunk_index=chunk_data['chunk_index'],
                        content=chunk_data['content']
                    )
                    chunk.set_file_info(chunk_data['metadata'])
                    db.session.add(chunk)

                # Update context with chunk count
                context.total_chunks = len(chunks_data)
                context.total_tokens = sum(len(chunk['content'].split()) for chunk in chunks_data)
                context.status = 'ready'
                context.progress = 100

                db.session.commit()

                return jsonify({
                    'files': uploaded_files,
                    'chunks_created': len(chunks_data),
                    'total_tokens': context.total_tokens,
                    'status': 'ready'
                }), 200

            except Exception as e:
                context.status = 'error'
                context.error_message = str(e)
                db.session.commit()
                return jsonify({'error': f'Processing failed: {str(e)}'}), 500

        return jsonify({'files': uploaded_files}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload/extract-zip', methods=['POST'])
@jwt_required()
def extract_zip():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No zip file provided'}), 400

        if 'context_id' not in request.form:
            return jsonify({'error': 'No context_id provided'}), 400

        zip_file = request.files['file']
        context_id = request.form['context_id']

        if zip_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # For now, just return a success message
        # In a real implementation, you would extract the zip file
        # and process its contents

        return jsonify({
            'message': 'Zip file extraction started',
            'files': [],
            'context_id': context_id
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload/supported-extensions', methods=['GET'])
def get_supported_extensions():
    extensions = {
        'documents': ['.pdf', '.docx', '.doc', '.txt', '.md'],
        'code': ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.go', '.rs'],
        'data': ['.csv', '.json', '.yaml', '.yml', '.xml'],
        'archives': ['.zip', '.tar', '.gz']
    }
    
    total_count = sum(len(exts) for exts in extensions.values())
    
    return jsonify({
        'extensions': extensions,
        'total_count': total_count
    }), 200

# Health check
@app.route('/api/database/test-connection', methods=['POST'])
@jwt_required()
def test_database_connection():
    """Test database connection and return available tables"""
    try:
        from services.database_service import DatabaseService

        data = request.get_json()
        db_type = data.get('db_type')
        connection_string = data.get('connection_string')

        if not db_type or not connection_string:
            return jsonify({'error': 'Database type and connection string required'}), 400

        db_service = DatabaseService()
        result = db_service.test_connection(db_type, connection_string)

        if result['success']:
            # Get table list
            tables = db_service.get_table_list(db_type, connection_string)
            result['tables'] = tables

        return jsonify(result), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'version': 'local-dev'
    }), 200

# Register route blueprints
from routes.admin import admin_bp
from routes.auth import auth_bp
from routes.contexts import contexts_bp
from routes.enhanced_contexts import enhanced_contexts_bp
from routes.chat import chat_bp
from routes.upload import upload_bp
from routes.versions import versions_bp
from routes.tasks import tasks_bp
from routes.preferences import preferences_bp

app.register_blueprint(admin_bp, url_prefix='/api/admin')
app.register_blueprint(auth_bp, url_prefix='/api/auth')
app.register_blueprint(contexts_bp, url_prefix='/api/contexts')
app.register_blueprint(enhanced_contexts_bp, url_prefix='/api/contexts')
app.register_blueprint(chat_bp, url_prefix='/api/chat')
app.register_blueprint(upload_bp, url_prefix='/api/upload')
app.register_blueprint(versions_bp, url_prefix='/api')
app.register_blueprint(tasks_bp, url_prefix='/api')
app.register_blueprint(preferences_bp)

# Error reporting endpoint
@app.route('/api/errors/report', methods=['POST'])
def report_error():
    """Simple error reporting endpoint"""
    try:
        error_data = request.get_json()
        app_logger.warning(f"Frontend error reported: {error_data}")
        return jsonify({'status': 'error_reported', 'timestamp': datetime.now(timezone.utc).isoformat()}), 200
    except Exception as e:
        app_logger.error(f"Error handling error report: {e}")
        return jsonify({'error': 'Failed to process error report'}), 500

# Debug endpoint to check JWT token
@app.route('/api/debug/token', methods=['GET'])
@jwt_required()
def debug_token():
    try:
        user_id = get_current_user_id()
        return jsonify({
            'status': 'Token is valid',
            'user_id': user_id,
            'timestamp': datetime.now(timezone.utc).isoformat()
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Debug endpoint to check headers (no auth required)
@app.route('/api/debug/headers', methods=['GET'])
def debug_headers():
    headers = dict(request.headers)
    return jsonify({
        'headers': headers,
        'has_authorization': 'Authorization' in headers,
        'authorization_header': headers.get('Authorization', 'Not present')
    }), 200

# JWT Error handlers
@jwt.expired_token_loader
def expired_token_callback(jwt_header, jwt_payload):
    print(f"JWT expired: {jwt_payload}")
    return jsonify({'error': 'Token has expired'}), 401

@jwt.invalid_token_loader
def invalid_token_callback(error):
    print(f"JWT invalid: {error}")
    return jsonify({'error': 'Invalid token'}), 401

@jwt.unauthorized_loader
def missing_token_callback(error):
    print(f"JWT missing: {error}")
    print(f"Request headers: {dict(request.headers)}")
    return jsonify({'error': 'Authorization token is required'}), 401

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500


# Security headers
@app.after_request
def add_security_headers(response):
    """Add security headers to all responses"""
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response

# Add request logging
@app.before_request
def log_request():
    """Log incoming requests"""
    request.start_time = datetime.now()
    
@app.after_request  
def log_response(response):
    """Log outgoing responses with timing"""
    try:
        if hasattr(request, 'start_time'):
            duration = (datetime.now() - request.start_time).total_seconds()
            log_request_info(request, response, duration)
    except Exception as e:
        app_logger.error(f"Error logging request: {e}")
    return response

@app.route('/api/admin/make-admin', methods=['POST'])
@jwt_required()
def make_first_user_admin():
    """Make the first user an admin (development convenience)"""
    try:
        user_id = get_current_user_id()
        
        # Only allow if there are no admins yet
        admin_count = User.query.filter_by(is_admin=True).count()
        if admin_count > 0:
            return jsonify({'error': 'Admin already exists'}), 403
        
        user = db.session.get(User, user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        user.is_admin = True
        db.session.commit()
        
        return jsonify({
            'message': f'User {user.username} is now an admin',
            'user': user.to_dict()
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def migrate_database():
    """Add missing columns to existing database"""
    try:
        from sqlalchemy import inspect, text
        
        # Check if is_admin column exists
        inspector = inspect(db.engine)
        columns = [col['name'] for col in inspector.get_columns('users')]
        
        if 'is_admin' not in columns:
            print("Adding is_admin column to users table...")
            with db.engine.connect() as conn:
                conn.execute(text('ALTER TABLE users ADD COLUMN is_admin BOOLEAN DEFAULT 0'))
                conn.commit()
            print("[SUCCESS] Added is_admin column successfully")
        else:
            print("[SUCCESS] is_admin column already exists")
            
    except Exception as e:
        print(f"Migration error: {e}")
        print("[ERROR] Database migration failed. Please delete the ragchatbot.db file and restart to create a fresh database.")

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # Run database migration for existing databases
        migrate_database()
        
        print("Database tables created successfully")
        print(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
        
        # Start background task service
        from services.task_service import start_task_service
        start_task_service(num_workers=2)
        print("Background task service started with 2 workers")
        
        print("Starting Flask development server...")

    # Run without debug mode to avoid watchdog issues
    app.run(debug=False, host='0.0.0.0', port=5000)
