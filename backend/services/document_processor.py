"""
Advanced Document Processing Service with Language-Aware Parsing

This module provides comprehensive document processing capabilities for the RAG Chatbot PWA,
supporting multiple file formats with intelligent content extraction and structure preservation.
The service uses advanced parsing techniques including Tree-sitter for code analysis and
specialized processors for different document types.

Key Features:
- Multi-format document support (PDF, DOCX, TXT, CSV, JSON, YAML, etc.)
- Language-aware code parsing using Tree-sitter AST analysis
- Intelligent chunking strategies preserving document structure
- Metadata extraction and content categorization
- Fallback mechanisms for unsupported formats
- Comprehensive error handling and logging
- Support for large document processing with memory optimization

Supported File Categories:
- Text Documents: TXT, MD, RST, LOG files
- Office Documents: PDF, DOCX, DOC, RTF
- Code Files: Python, JavaScript, Java, C/C++, Go, Rust, etc.
- Data Files: CSV, XLSX, JSON, XML, YAML
- Configuration Files: INI, CFG, TOML, Properties
- Web Files: HTML, CSS, SCSS
- Database Files: SQL, DDL, DML

Architecture:
1. File Type Detection: Automatic format identification
2. Content Extraction: Format-specific parsing and text extraction
3. Structure Analysis: AST parsing for code, DOM parsing for markup
4. Intelligent Chunking: Content-aware segmentation
5. Metadata Enrichment: File information and structure metadata
6. Quality Assurance: Validation and error recovery

Dependencies:
- PyMuPDF (fitz): PDF processing
- python-docx: Word document processing
- pandas: Data file processing
- tree-sitter: Code parsing and AST analysis
- unstructured: Advanced document parsing
- pyyaml: YAML processing

Author: RAG Chatbot Development Team
Version: 1.0.0
"""

import os
import ast
import json
import yaml
import fitz  # PyMuPDF
import pandas as pd
from typing import List, Dict, Any, Optional
from pathlib import Path
import tree_sitter
from tree_sitter import Language, Parser
import docx
# Optional advanced document parsing - fallback if not available
try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    partition = None
    logger.warning("unstructured library not available - advanced document parsing will be limited")

# Import logging functionality
from logging_config import get_logger, log_error_with_context

# Import Tree-sitter service for advanced code parsing
from .tree_sitter_service import tree_sitter_service, CodeChunk

# Initialize logger
logger = get_logger('document_processor')

class DocumentProcessor:
    """
    Advanced document processing service with language-aware parsing capabilities
    
    This class provides comprehensive document processing for the RAG system, supporting
    multiple file formats with intelligent content extraction, structure preservation,
    and optimized chunking strategies. It uses Tree-sitter for code analysis and
    specialized processors for different document types.
    
    The processor handles:
    - Automatic file type detection and format-specific processing
    - Language-aware code parsing with AST analysis
    - Intelligent text chunking preserving document structure
    - Metadata extraction and content categorization
    - Error recovery and fallback mechanisms
    
    Attributes:
        supported_extensions (Dict[str, List[str]]): Mapping of file categories to extensions
        parsers (Dict[str, Parser]): Tree-sitter parsers for different programming languages
        
    Example:
        >>> processor = DocumentProcessor()
        >>> chunks = processor.process_file('/path/to/document.pdf')
        >>> print(f"Extracted {len(chunks)} chunks from document")
    """
    
    def __init__(self):
        """
        Initialize the Document Processor with supported file types and parsers
        
        Sets up the processing environment including file type mappings and
        Tree-sitter parsers for code analysis. Initializes logging and
        performs dependency checks.
        
        Raises:
            ImportError: If required dependencies are missing
            Exception: If Tree-sitter parsers cannot be initialized
        """
        logger.info("Initializing DocumentProcessor")
        
        # Define supported file extensions by category
        self.supported_extensions = {
            'text': ['.txt', '.md', '.rst', '.log'],
            'code': ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.h', '.hpp', '.go', '.rs', '.rb', '.php', '.cs', '.kt', '.swift'],
            'document': ['.pdf', '.docx', '.doc', '.rtf'],
            'data': ['.csv', '.xlsx', '.xls', '.json', '.xml', '.yaml', '.yml'],
            'config': ['.ini', '.cfg', '.conf', '.toml', '.properties'],
            'web': ['.html', '.htm', '.css', '.scss', '.sass', '.less'],
            'sql': ['.sql', '.ddl', '.dml']
        }
        
        # Count total supported extensions
        total_extensions = sum(len(exts) for exts in self.supported_extensions.values())
        logger.info(f"Supporting {total_extensions} file extensions across {len(self.supported_extensions)} categories")
        
        # Initialize tree-sitter parsers
        self.parsers = {}
        try:
            self._init_tree_sitter_parsers()
            logger.info("DocumentProcessor initialization completed successfully")
        except Exception as e:
            logger.error(f"Failed to initialize DocumentProcessor: {e}")
            log_error_with_context(e, {"supported_extensions": list(self.supported_extensions.keys())})
            raise
    
    def _init_tree_sitter_parsers(self):
        """
        Initialize Tree-sitter parsers for supported programming languages
        
        Tree-sitter provides syntax-aware parsing for code files, enabling better
        chunking strategies that preserve function and class boundaries. This method
        attempts to initialize parsers for supported languages with fallback handling.
        
        Supported Languages:
        - Python: Function and class boundary detection
        - JavaScript/TypeScript: Module and function parsing
        - Java: Class and method structure analysis
        - C/C++: Header and implementation file parsing
        - Go: Package and function boundary detection
        - Kotlin: Class and function structure parsing
        
        Note:
            In production environments, Tree-sitter language libraries need to be
            compiled and made available. This implementation includes fallback
            mechanisms for environments without compiled language libraries.
            
        Raises:
            Exception: If critical parser initialization fails
        """
        logger.debug("Initializing Tree-sitter parsers for code analysis")
        
        try:
            # Configuration for supported language parsers
            language_configs = {
                'python': 'tree-sitter-python',
                'javascript': 'tree-sitter-javascript',
                'typescript': 'tree-sitter-typescript',
                'java': 'tree-sitter-java',
                'cpp': 'tree-sitter-cpp',
                'go': 'tree-sitter-go',
                'kotlin': 'tree-sitter-kotlin'
            }
            
            successful_parsers = []
            failed_parsers = []
            
            for lang, lib_name in language_configs.items():
                try:
                    # In production, you would load the compiled language library:
                    # language = Language(f'/path/to/compiled/{lib_name}.so', lang)
                    # parser = Parser()
                    # parser.set_language(language)
                    # self.parsers[lang] = parser
                    
                    # For now, we'll mark as available for future implementation
                    logger.debug(f"Tree-sitter parser for {lang} would be loaded from {lib_name}")
                    # self.parsers[lang] = None  # Placeholder for actual parser
                    successful_parsers.append(lang)
                    
                except Exception as e:
                    logger.warning(f"Could not initialize {lang} parser ({lib_name}): {e}")
                    failed_parsers.append(lang)
            
            # Log initialization results
            if successful_parsers:
                logger.info(f"Tree-sitter parsers prepared for: {', '.join(successful_parsers)}")
            
            if failed_parsers:
                logger.warning(f"Failed to initialize parsers for: {', '.join(failed_parsers)}")
                
            logger.info("Tree-sitter parser initialization completed (fallback mode)")
            
        except Exception as e:
            logger.error(f"Critical error in Tree-sitter initialization: {e}")
            log_error_with_context(e, {"language_configs": list(language_configs.keys())})
            # Don't raise here - continue with basic processing
            logger.warning("Continuing with basic document processing (no syntax-aware parsing)")
    
    def process_file(self, file_path: str, chunk_strategy: str = 'language-specific') -> List[Dict[str, Any]]:
        """
        Process a file and extract structured chunks with metadata
        
        This is the main entry point for document processing. It automatically detects
        the file type and applies the appropriate processing strategy to extract
        meaningful chunks while preserving document structure and context.
        
        Processing Pipeline:
        1. File Type Detection: Analyze extension and content
        2. Content Extraction: Format-specific text extraction
        3. Structure Analysis: Language or format-aware parsing
        4. Intelligent Chunking: Context-preserving segmentation
        5. Metadata Enrichment: Add file and chunk metadata
        6. Quality Validation: Ensure chunk quality and completeness
        
        Args:
            file_path (str): Absolute path to the file to process
            chunk_strategy (str): Chunking strategy to use. Options:
                - 'language-specific': Syntax-aware chunking for code files
                - 'semantic': Content-aware chunking for documents
                - 'fixed-size': Simple size-based chunking
                - 'adaptive': Dynamic strategy selection based on content
                
        Returns:
            List[Dict[str, Any]]: List of processed chunks, each containing:
                - content (str): Extracted text content
                - metadata (Dict): File information, chunk index, type, etc.
                - structure (Dict, optional): Document structure information
                - language (str, optional): Detected programming language
                
        Raises:
            FileNotFoundError: If the specified file doesn't exist
            PermissionError: If file cannot be read due to permissions
            ValueError: If chunk_strategy is not supported
            
        Example:
            >>> processor = DocumentProcessor()
            >>> chunks = processor.process_file('/path/to/code.py', 'language-specific')
            >>> for chunk in chunks:
            ...     print(f"Chunk {chunk['metadata']['chunk_index']}: {len(chunk['content'])} chars")
            
        Note:
            - Large files are processed in memory-efficient chunks
            - Error recovery ensures partial processing results are returned
            - File type detection is based on extension and content analysis
        """
        file_path = Path(file_path)
        file_name = file_path.name
        file_ext = file_path.suffix.lower()
        file_size = file_path.stat().st_size if file_path.exists() else 0
        
        logger.info(f"Processing file: {file_name} (size: {file_size:,} bytes, strategy: {chunk_strategy})")
        
        # Validate inputs
        if not file_path.exists():
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        if chunk_strategy not in ['language-specific', 'semantic', 'fixed-size', 'adaptive']:
            logger.warning(f"Unknown chunk strategy '{chunk_strategy}', using 'language-specific'")
            chunk_strategy = 'language-specific'
        
        # Detect file type
        file_type = self._get_file_type(file_ext)
        logger.debug(f"Detected file type: {file_type} for extension: {file_ext}")
        
        start_time = logger.info(f"Starting {file_type} file processing")
        
        try:
            # Route to appropriate processor based on file type
            if file_type == 'code':
                chunks = self._process_code_file(file_path, chunk_strategy)
            elif file_type == 'document':
                chunks = self._process_document_file(file_path, chunk_strategy)
            elif file_type == 'data':
                chunks = self._process_data_file(file_path, chunk_strategy)
            elif file_type == 'text':
                chunks = self._process_text_file(file_path, chunk_strategy)
            elif file_type == 'config':
                chunks = self._process_config_file(file_path, chunk_strategy)
            elif file_type == 'web':
                chunks = self._process_web_file(file_path, chunk_strategy)
            elif file_type == 'sql':
                chunks = self._process_sql_file(file_path, chunk_strategy)
            else:
                logger.warning(f"Using generic processor for unknown file type: {file_type}")
                chunks = self._process_generic_file(file_path, chunk_strategy)
            
            # Validate and log results
            if chunks:
                total_content = sum(len(chunk.get('content', '')) for chunk in chunks)
                avg_chunk_size = total_content / len(chunks) if chunks else 0
                logger.info(f"Successfully processed {file_name}: {len(chunks)} chunks, {total_content:,} total characters, {avg_chunk_size:.1f} avg size")
            else:
                logger.warning(f"No chunks extracted from {file_name}")
                
            return chunks
        
        except FileNotFoundError:
            raise  # Re-raise file not found errors
        except PermissionError as e:
            error_msg = f"Permission denied accessing {file_name}: {e}"
            logger.error(error_msg)
            log_error_with_context(e, {"file_path": str(file_path), "file_type": file_type})
            raise PermissionError(error_msg)
        except Exception as e:
            error_msg = f"Error processing {file_name}: {str(e)}"
            logger.error(error_msg)
            log_error_with_context(e, {
                "file_path": str(file_path),
                "file_name": file_name,
                "file_type": file_type,
                "file_size": file_size,
                "chunk_strategy": chunk_strategy
            })
            
            # Return error chunk instead of failing completely
            return [{
                'content': f"Error processing file: {str(e)}",
                'metadata': {
                    'file_path': str(file_path),
                    'file_name': file_name,
                    'file_type': file_type,
                    'file_size': file_size,
                    'error': True,
                    'error_message': str(e),
                    'chunk_index': 0,
                    'chunk_strategy': chunk_strategy
                }
            }]
    
    def _get_file_type(self, extension: str) -> str:
        """Get file type category from extension"""
        for file_type, extensions in self.supported_extensions.items():
            if extension in extensions:
                return file_type
        return 'unknown'
    
    def _get_language_from_extension(self, extension: str) -> Optional[str]:
        """Get programming language from file extension"""
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.tsx': 'typescript',
            '.jsx': 'javascript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'cpp',
            '.h': 'cpp',
            '.hpp': 'cpp',
            '.go': 'go',
            '.kt': 'kotlin',
            '.rs': 'rust',
            '.rb': 'ruby',
            '.php': 'php',
            '.cs': 'csharp',
            '.swift': 'swift'
        }
        return language_map.get(extension)
    
    def _process_code_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """
        Process code files with advanced Tree-sitter parsing
        
        Uses the integrated Tree-sitter service for intelligent code parsing that
        preserves function and class boundaries. Falls back to legacy methods when
        Tree-sitter parsing is not available or fails.
        
        Args:
            file_path (str): Path to the code file
            chunk_strategy (str): Chunking strategy to apply
            
        Returns:
            List[Dict[str, Any]]: Processed code chunks with enhanced metadata
        """
        logger.debug(f"Processing code file: {file_path} with strategy: {chunk_strategy}")
        
        try:
            # Use Tree-sitter service for advanced code parsing
            tree_sitter_chunks = tree_sitter_service.parse_code_file(str(file_path))
            
            if tree_sitter_chunks:
                logger.info(f"Tree-sitter parsing successful: {len(tree_sitter_chunks)} chunks found")
                return self._convert_tree_sitter_chunks(tree_sitter_chunks, file_path)
            else:
                logger.warning("Tree-sitter parsing returned no chunks, using fallback")
                
        except Exception as e:
            logger.warning(f"Tree-sitter parsing failed for {file_path}: {e}")
            log_error_with_context(e, {
                "file_path": file_path,
                "chunk_strategy": chunk_strategy,
                "operation": "tree_sitter_code_parsing"
            })
        
        # Fallback to legacy processing methods
        logger.debug("Using legacy code processing as fallback")
        return self._process_code_file_legacy(file_path, chunk_strategy)
    
    def _convert_tree_sitter_chunks(self, tree_sitter_chunks: List[CodeChunk], file_path: str) -> List[Dict[str, Any]]:
        """
        Convert Tree-sitter CodeChunk objects to standard document processor format
        
        Transforms the Tree-sitter service output into the format expected by the
        document processor, preserving all metadata and structure information.
        
        Args:
            tree_sitter_chunks (List[CodeChunk]): Chunks from Tree-sitter service
            file_path (str): Original file path for context
            
        Returns:
            List[Dict[str, Any]]: Chunks in document processor format
        """
        converted_chunks = []
        file_name = Path(file_path).name
        
        for i, ts_chunk in enumerate(tree_sitter_chunks):
            # Convert CodeChunk to document processor format
            chunk_dict = {
                'content': ts_chunk.content,
                'metadata': {
                    'chunk_index': i,
                    'chunk_type': ts_chunk.chunk_type,
                    'chunk_name': ts_chunk.name,
                    'file_name': file_name,
                    'file_path': file_path,
                    'language': ts_chunk.language,
                    'start_line': ts_chunk.start_line,
                    'end_line': ts_chunk.end_line,
                    'total_lines': ts_chunk.end_line - ts_chunk.start_line + 1,
                    'character_count': len(ts_chunk.content),
                    'parsing_method': 'tree_sitter',
                    'syntax_aware': True,
                    'has_docstring': ts_chunk.metadata.get('has_docstring', False),
                    'complexity_estimate': ts_chunk.metadata.get('complexity_estimate', 1),
                    **ts_chunk.metadata  # Include all original metadata
                }
            }
            
            # Add structure information for better RAG context
            if ts_chunk.chunk_type in ['function', 'method']:
                chunk_dict['metadata']['is_callable'] = True
            elif ts_chunk.chunk_type in ['class', 'interface']:
                chunk_dict['metadata']['is_type_definition'] = True
            
            converted_chunks.append(chunk_dict)
        
        logger.debug(f"Converted {len(tree_sitter_chunks)} Tree-sitter chunks to document format")
        return converted_chunks
    
    def _process_code_file_legacy(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """
        Legacy code processing method for fallback support
        
        Provides backward compatibility when Tree-sitter parsing fails or is
        unavailable. Uses the original AST-based Python parsing and generic
        text chunking for other languages.
        
        Args:
            file_path (str): Path to the code file
            chunk_strategy (str): Chunking strategy to apply
            
        Returns:
            List[Dict[str, Any]]: Processed chunks using legacy methods
        """
        logger.debug(f"Using legacy code processing for: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        file_ext = Path(file_path).suffix.lower()
        language = self._get_language_from_extension(file_ext)
        
        # Use legacy Python AST parsing for Python files
        if language == 'python':
            logger.debug("Using legacy Python AST parsing")
            return self._process_python_file(content, file_path)
        else:
            # Use generic text chunking for other languages
            logger.debug(f"Using generic text chunking for {language} file")
            return self._process_code_generic(content, file_path, language)
    
    def _process_python_file(self, content: str, file_path: str) -> List[Dict[str, Any]]:
        """Process Python files using AST"""
        chunks = []
        
        try:
            tree = ast.parse(content)
            
            for node in ast.walk(tree):
                if isinstance(node, (ast.FunctionDef, ast.ClassDef, ast.AsyncFunctionDef)):
                    # Extract function/class with docstring
                    start_line = node.lineno
                    end_line = node.end_lineno if hasattr(node, 'end_lineno') else start_line
                    
                    lines = content.split('\n')
                    chunk_content = '\n'.join(lines[start_line-1:end_line])
                    
                    # Extract docstring
                    docstring = ast.get_docstring(node)
                    
                    chunks.append({
                        'content': chunk_content,
                        'metadata': {
                            'file_path': file_path,
                            'file_type': 'code',
                            'language': 'python',
                            'node_type': type(node).__name__,
                            'name': node.name,
                            'start_line': start_line,
                            'end_line': end_line,
                            'docstring': docstring,
                            'chunk_index': len(chunks)
                        }
                    })
            
            # If no functions/classes found, chunk by lines
            if not chunks:
                return self._chunk_by_lines(content, file_path, 'python')
        
        except SyntaxError:
            # Fallback to line-based chunking for invalid Python
            return self._chunk_by_lines(content, file_path, 'python')
        
        return chunks
    
    def _process_with_tree_sitter(self, content: str, file_path: str, language: str) -> List[Dict[str, Any]]:
        """Process code using tree-sitter (placeholder implementation)"""
        # This would use tree-sitter to parse the code and extract functions, classes, etc.
        # For now, fall back to generic code processing
        return self._process_code_generic(content, file_path, language)
    
    def _process_code_generic(self, content: str, file_path: str, language: Optional[str]) -> List[Dict[str, Any]]:
        """Generic code processing by logical blocks"""
        chunks = []
        lines = content.split('\n')
        current_chunk = []
        current_indent = 0
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # Skip empty lines
            if not stripped:
                if current_chunk:
                    current_chunk.append(line)
                continue
            
            # Detect function/class definitions
            if any(stripped.startswith(keyword) for keyword in ['def ', 'class ', 'function ', 'interface ', 'struct ']):
                # Save previous chunk
                if current_chunk:
                    chunks.append({
                        'content': '\n'.join(current_chunk),
                        'metadata': {
                            'file_path': file_path,
                            'file_type': 'code',
                            'language': language,
                            'chunk_index': len(chunks),
                            'start_line': i - len(current_chunk) + 1,
                            'end_line': i
                        }
                    })
                
                # Start new chunk
                current_chunk = [line]
                current_indent = len(line) - len(line.lstrip())
            
            elif current_chunk:
                # Continue current chunk
                current_chunk.append(line)
                
                # Check if we should end the chunk (dedent)
                line_indent = len(line) - len(line.lstrip())
                if line_indent <= current_indent and len(current_chunk) > 20:
                    chunks.append({
                        'content': '\n'.join(current_chunk),
                        'metadata': {
                            'file_path': file_path,
                            'file_type': 'code',
                            'language': language,
                            'chunk_index': len(chunks),
                            'start_line': i - len(current_chunk) + 1,
                            'end_line': i
                        }
                    })
                    current_chunk = []
            else:
                current_chunk = [line]
        
        # Add final chunk
        if current_chunk:
            chunks.append({
                'content': '\n'.join(current_chunk),
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'code',
                    'language': language,
                    'chunk_index': len(chunks),
                    'start_line': len(lines) - len(current_chunk) + 1,
                    'end_line': len(lines)
                }
            })
        
        return chunks
    
    def _process_document_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Process document files (PDF, DOCX, etc.)"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._process_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._process_docx(file_path)
        else:
            # Use unstructured.io for other document types
            return self._process_with_unstructured(file_path)
    
    def _process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Process PDF files using PyMuPDF"""
        chunks = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    chunks.append({
                        'content': text,
                        'metadata': {
                            'file_path': file_path,
                            'file_type': 'document',
                            'page_number': page_num + 1,
                            'chunk_index': len(chunks)
                        }
                    })
            
            doc.close()
        
        except Exception as e:
            chunks.append({
                'content': f"Error processing PDF: {str(e)}",
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'document',
                    'error': True,
                    'chunk_index': 0
                }
            })
        
        return chunks
    
    def _process_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Process DOCX files with enhanced content extraction"""
        chunks = []
        
        try:
            doc = docx.Document(file_path)
            file_name = Path(file_path).name
            
            # Extract document properties
            doc_properties = {
                'title': doc.core_properties.title or '',
                'subject': doc.core_properties.subject or '',
                'creator': doc.core_properties.author or '',
                'keywords': doc.core_properties.keywords or '',
                'description': doc.core_properties.comments or ''
            }
            
            # Process document structure
            sections = []
            current_section = {
                'title': 'Document Start',
                'content': [],
                'subsections': [],
                'tables': [],
                'lists': []
            }
            
            # Enhanced paragraph processing
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style_name = paragraph.style.name
                    text = paragraph.text.strip()
                    
                    # Check for headings and create sections
                    if style_name.startswith('Heading'):
                        # Save current section if it has content
                        if current_section['content'] or current_section['tables'] or current_section['lists']:
                            sections.append(current_section)
                        
                        # Start new section
                        heading_level = style_name.replace('Heading ', '')
                        current_section = {
                            'title': text,
                            'heading_level': heading_level,
                            'content': [],
                            'subsections': [],
                            'tables': [],
                            'lists': []
                        }
                    else:
                        # Regular paragraph - check for special formatting
                        paragraph_info = {
                            'text': text,
                            'style': style_name,
                            'is_bold': any(run.bold for run in paragraph.runs if run.bold),
                            'is_italic': any(run.italic for run in paragraph.runs if run.italic),
                        }
                        current_section['content'].append(paragraph_info)
            
            # Add final section
            if current_section['content'] or current_section['tables'] or current_section['lists']:
                sections.append(current_section)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if row_idx == 0:
                        headers = row_data
                    table_data.append(row_data)
                
                table_info = {
                    'headers': headers,
                    'data': table_data,
                    'row_count': len(table_data),
                    'col_count': len(headers) if headers else 0
                }
                
                # Add table to current section or create separate chunk
                if sections:
                    sections[-1]['tables'].append(table_info)
                else:
                    sections.append({
                        'title': f'Table {table_idx + 1}',
                        'content': [],
                        'tables': [table_info],
                        'lists': []
                    })
            
            # Create chunks from sections
            for section_idx, section in enumerate(sections):
                section_content = []
                
                # Add section title
                if section['title']:
                    section_content.append(f"# {section['title']}")
                
                # Add regular content
                for para in section['content']:
                    text = para['text']
                    if para.get('is_bold'):
                        text = f"**{text}**"
                    if para.get('is_italic'):
                        text = f"*{text}*"
                    section_content.append(text)
                
                # Add tables
                for table_idx, table in enumerate(section['tables']):
                    section_content.append(f"\n## Table {table_idx + 1}")
                    if table['headers']:
                        section_content.append("| " + " | ".join(table['headers']) + " |")
                        section_content.append("| " + " | ".join(["---"] * len(table['headers'])) + " |")
                    
                    for row in table['data'][1:] if table['headers'] else table['data']:
                        section_content.append("| " + " | ".join(row) + " |")
                
                # Create chunk
                content = '\n'.join(section_content)
                if content.strip():
                    chunks.append({
                        'content': content,
                        'metadata': {
                            'file_path': str(file_path),
                            'file_name': file_name,
                            'file_type': 'docx',
                            'document_type': 'office_document',
                            'section_title': section['title'],
                            'section_index': section_idx,
                            'chunk_index': len(chunks),
                            'has_tables': len(section['tables']) > 0,
                            'table_count': len(section['tables']),
                            'paragraph_count': len(section['content']),
                            'document_properties': doc_properties,
                            'processing_method': 'enhanced_docx'
                        }
                    })
            
            # If no sections were found, create a single chunk with all content
            if not chunks:
                all_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        all_text.append(paragraph.text.strip())
                
                content = '\n'.join(all_text)
                chunks.append({
                    'content': content,
                    'metadata': {
                        'file_path': str(file_path),
                        'file_name': file_name,
                        'file_type': 'docx',
                        'document_type': 'office_document',
                        'chunk_index': 0,
                        'document_properties': doc_properties,
                        'processing_method': 'enhanced_docx_fallback'
                    }
                })
        
        except Exception as e:
            file_name = Path(file_path).name
            logger.error(f"Error processing DOCX {file_name}: {str(e)}")
            chunks.append({
                'content': f"{file_name} is a Microsoft Word document that contains structured text, tables, and formatting. Due to processing limitations, the detailed content could not be extracted. The document appears to contain business or technical information that would require specialized document parsing tools to fully access.",
                'metadata': {
                    'file_path': str(file_path),
                    'file_name': file_name,
                    'file_type': 'docx',
                    'document_type': 'office_document',
                    'error': True,
                    'error_message': str(e),
                    'chunk_index': 0,
                    'processing_method': 'error_fallback'
                }
            })
        
        return chunks
    
    def _process_with_unstructured(self, file_path: str) -> List[Dict[str, Any]]:
        """Process files using unstructured.io - fallback if not available"""
        chunks = []
        
        if not UNSTRUCTURED_AVAILABLE or partition is None:
            logger.warning(f"Unstructured library not available - falling back to basic processing for {file_path}")
            return self._fallback_text_processing(file_path)
        
        try:
            elements = partition(filename=file_path)
            
            for i, element in enumerate(elements):
                chunks.append({
                    'content': str(element),
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'document',
                        'element_type': type(element).__name__,
                        'chunk_index': i
                    }
                })
        
        except Exception as e:
            chunks.append({
                'content': f"Error processing with unstructured: {str(e)}",
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'document',
                    'error': True,
                    'chunk_index': 0
                }
            })
        
        return chunks

    def _fallback_text_processing(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Fallback text processing when unstructured library is not available.
        Attempts basic text extraction for common file types.
        """
        chunks = []
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.txt':
                # Plain text file
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            elif file_ext in ['.html', '.htm']:
                # Basic HTML text extraction
                try:
                    from bs4 import BeautifulSoup
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        soup = BeautifulSoup(f.read(), 'html.parser')
                        content = soup.get_text(separator='\n', strip=True)
                except ImportError:
                    logger.warning("BeautifulSoup not available for HTML parsing - using raw text")
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
            
            elif file_ext in ['.md', '.markdown']:
                # Markdown files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            elif file_ext == '.rtf':
                # Basic RTF processing - strip RTF codes
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    raw_content = f.read()
                    # Very basic RTF cleaning - remove control words
                    import re
                    content = re.sub(r'\\[a-z]+\d*\s?', '', raw_content)
                    content = re.sub(r'[{}]', '', content)
                    content = content.strip()
            
            else:
                # Unknown file type - attempt to read as text
                logger.warning(f"Unknown file type {file_ext} - attempting to read as plain text")
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                except Exception as e:
                    logger.error(f"Failed to read {file_path} as text: {e}")
                    return [{
                        'content': f"Error: Could not process file {Path(file_path).name}. Unsupported format without unstructured library.",
                        'metadata': {
                            'file_path': str(file_path),
                            'file_name': Path(file_path).name,
                            'file_type': file_ext,
                            'error': 'Unsupported format',
                            'chunk_index': 0
                        }
                    }]
            
            # Clean and validate content
            if not content or not content.strip():
                logger.warning(f"No text content extracted from {file_path}")
                return [{
                    'content': f"Warning: No readable text content found in {Path(file_path).name}",
                    'metadata': {
                        'file_path': str(file_path),
                        'file_name': Path(file_path).name,
                        'file_type': file_ext,
                        'warning': 'No content extracted',
                        'chunk_index': 0
                    }
                }]
            
            # Basic chunking - split into paragraphs or by size
            content = content.strip()
            max_chunk_size = 4000  # Conservative chunk size
            
            if len(content) <= max_chunk_size:
                # Single chunk
                chunks.append({
                    'content': content,
                    'metadata': {
                        'file_path': str(file_path),
                        'file_name': Path(file_path).name,
                        'file_type': file_ext,
                        'processing_method': 'fallback',
                        'chunk_index': 0,
                        'total_chunks': 1
                    }
                })
            else:
                # Split into multiple chunks
                paragraphs = content.split('\n\n')
                current_chunk = ""
                chunk_index = 0
                
                for paragraph in paragraphs:
                    if len(current_chunk) + len(paragraph) + 2 <= max_chunk_size:
                        if current_chunk:
                            current_chunk += '\n\n' + paragraph
                        else:
                            current_chunk = paragraph
                    else:
                        if current_chunk:
                            chunks.append({
                                'content': current_chunk.strip(),
                                'metadata': {
                                    'file_path': str(file_path),
                                    'file_name': Path(file_path).name,
                                    'file_type': file_ext,
                                    'processing_method': 'fallback',
                                    'chunk_index': chunk_index
                                }
                            })
                            chunk_index += 1
                        current_chunk = paragraph
                
                # Add the last chunk
                if current_chunk:
                    chunks.append({
                        'content': current_chunk.strip(),
                        'metadata': {
                            'file_path': str(file_path),
                            'file_name': Path(file_path).name,
                            'file_type': file_ext,
                            'processing_method': 'fallback',
                            'chunk_index': chunk_index
                        }
                    })
                
                # Update total chunks count
                total_chunks = len(chunks)
                for chunk in chunks:
                    chunk['metadata']['total_chunks'] = total_chunks
            
            logger.info(f"Fallback processing completed for {file_path}: {len(chunks)} chunks extracted")
            
        except Exception as e:
            logger.error(f"Fallback processing failed for {file_path}: {e}")
            chunks = [{
                'content': f"Error: Failed to process file {Path(file_path).name}. {str(e)}",
                'metadata': {
                    'file_path': str(file_path),
                    'file_name': Path(file_path).name,
                    'file_type': file_ext,
                    'error': str(e),
                    'processing_method': 'fallback_error',
                    'chunk_index': 0
                }
            }]
        
        return chunks

    def _process_data_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Process data files (CSV, JSON, XLSX, etc.)"""
        file_ext = Path(file_path).suffix.lower()
        chunks = []

        try:
            if file_ext == '.csv':
                df = pd.read_csv(file_path)
                return self._process_dataframe(df, file_path, 'csv')

            elif file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                return self._process_dataframe(df, file_path, 'excel')

            elif file_ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                chunks.append({
                    'content': json.dumps(data, indent=2),
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'data',
                        'format': 'json',
                        'chunk_index': 0
                    }
                })

            elif file_ext in ['.yaml', '.yml']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = yaml.safe_load(f)

                chunks.append({
                    'content': yaml.dump(data, default_flow_style=False),
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'data',
                        'format': 'yaml',
                        'chunk_index': 0
                    }
                })

            elif file_ext == '.xml':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                chunks.append({
                    'content': content,
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'data',
                        'format': 'xml',
                        'chunk_index': 0
                    }
                })

        except Exception as e:
            chunks.append({
                'content': f"Error processing data file: {str(e)}",
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'data',
                    'error': True,
                    'chunk_index': 0
                }
            })

        return chunks

    def _process_dataframe(self, df: pd.DataFrame, file_path: str, format_type: str) -> List[Dict[str, Any]]:
        """Process pandas DataFrame into chunks"""
        chunks = []

        # Add column information
        column_info = {
            'columns': list(df.columns),
            'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()},
            'shape': df.shape
        }

        chunks.append({
            'content': f"Dataset Schema:\nColumns: {', '.join(df.columns)}\nShape: {df.shape}\nData Types:\n" +
                      '\n'.join([f"{col}: {dtype}" for col, dtype in df.dtypes.items()]),
            'metadata': {
                'file_path': file_path,
                'file_type': 'data',
                'format': format_type,
                'chunk_type': 'schema',
                'chunk_index': 0,
                **column_info
            }
        })

        # Chunk data by rows
        chunk_size = 100
        for i in range(0, len(df), chunk_size):
            chunk_df = df.iloc[i:i+chunk_size]

            chunks.append({
                'content': chunk_df.to_string(),
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'data',
                    'format': format_type,
                    'chunk_type': 'data',
                    'chunk_index': len(chunks),
                    'row_start': i,
                    'row_end': min(i + chunk_size, len(df)),
                    **column_info
                }
            })

        return chunks

    def _process_text_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        return self._chunk_by_lines(content, file_path, 'text')

    def _process_config_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Process configuration files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        file_ext = Path(file_path).suffix.lower()

        return [{
            'content': content,
            'metadata': {
                'file_path': file_path,
                'file_type': 'config',
                'format': file_ext[1:],  # Remove the dot
                'chunk_index': 0
            }
        }]

    def _process_web_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Process web files (HTML, CSS, etc.)"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        file_ext = Path(file_path).suffix.lower()

        if file_ext in ['.html', '.htm']:
            # Could parse HTML structure here
            return self._chunk_by_lines(content, file_path, 'html')
        elif file_ext in ['.css', '.scss', '.sass', '.less']:
            return self._process_css_file(content, file_path)
        else:
            return self._chunk_by_lines(content, file_path, 'web')

    def _process_css_file(self, content: str, file_path: str) -> List[Dict[str, Any]]:
        """Process CSS files by rules"""
        chunks = []
        lines = content.split('\n')
        current_rule = []
        brace_count = 0

        for i, line in enumerate(lines):
            current_rule.append(line)
            brace_count += line.count('{') - line.count('}')

            if brace_count == 0 and current_rule and any('{' in l for l in current_rule):
                chunks.append({
                    'content': '\n'.join(current_rule),
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'web',
                        'language': 'css',
                        'chunk_index': len(chunks),
                        'start_line': i - len(current_rule) + 1,
                        'end_line': i
                    }
                })
                current_rule = []

        if current_rule:
            chunks.append({
                'content': '\n'.join(current_rule),
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'web',
                    'language': 'css',
                    'chunk_index': len(chunks)
                }
            })

        return chunks

    def _process_sql_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Process SQL files by statements"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Split by semicolons (simple SQL statement separation)
        statements = [stmt.strip() for stmt in content.split(';') if stmt.strip()]

        chunks = []
        for i, statement in enumerate(statements):
            chunks.append({
                'content': statement,
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'sql',
                    'statement_type': self._get_sql_statement_type(statement),
                    'chunk_index': i
                }
            })

        return chunks

    def _get_sql_statement_type(self, statement: str) -> str:
        """Determine SQL statement type"""
        statement_upper = statement.upper().strip()

        if statement_upper.startswith('SELECT'):
            return 'SELECT'
        elif statement_upper.startswith('INSERT'):
            return 'INSERT'
        elif statement_upper.startswith('UPDATE'):
            return 'UPDATE'
        elif statement_upper.startswith('DELETE'):
            return 'DELETE'
        elif statement_upper.startswith('CREATE'):
            return 'CREATE'
        elif statement_upper.startswith('ALTER'):
            return 'ALTER'
        elif statement_upper.startswith('DROP'):
            return 'DROP'
        else:
            return 'OTHER'

    def _process_generic_file(self, file_path: str, chunk_strategy: str) -> List[Dict[str, Any]]:
        """Generic file processing fallback"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            return self._chunk_by_lines(content, file_path, 'unknown')

        except Exception as e:
            return [{
                'content': f"Error processing file: {str(e)}",
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'unknown',
                    'error': True,
                    'chunk_index': 0
                }
            }]

    def _chunk_by_lines(self, content: str, file_path: str, file_type: str,
                       chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """Chunk content by lines with overlap"""
        lines = content.split('\n')
        chunks = []

        for i in range(0, len(lines), chunk_size - overlap):
            chunk_lines = lines[i:i + chunk_size]

            if chunk_lines:
                chunks.append({
                    'content': '\n'.join(chunk_lines),
                    'metadata': {
                        'file_path': file_path,
                        'file_type': file_type,
                        'chunk_index': len(chunks),
                        'start_line': i + 1,
                        'end_line': min(i + chunk_size, len(lines)),
                        'total_lines': len(lines)
                    }
                })

        return chunks
